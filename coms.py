from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_data_dictionary():
    doc = Document()
    
    # ===== DOCUMENT SETUP =====
    # Set margins for better table fitting
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
    
    # Title Page
    title = doc.add_heading('GOVHRPay - Data Dictionary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(24)
    
    subtitle = doc.add_paragraph('Complete Database Schema Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True
    
    doc.add_paragraph('\n\n')
    
    info = doc.add_paragraph('System: GOV HR & Payroll Management\n'
                           'Version: 1.0\n'
                           'Generated: April 11, 2026\n'
                           'Database: PostgreSQL/MySQL via SQLAlchemy')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents placeholder
    toc = doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Users (users) .................................... 3\n'
                     '2. Employees (employee) ............................. 4\n'
                     '3. Payroll Period (payroll_period) .................. 5\n'
                     '4. Payroll (payroll) ................................ 6\n'
                     '5. Payroll Deduction (payroll_deduction) ............ 7\n'
                     '6. Payslip (payslip) ................................ 8\n'
                     '7. Deduction (deduction) ............................ 9\n'
                     '8. Employee Deduction (employee_deductions) ......... 10\n'
                     '9. Deduction Bracket (deduction_bracket) ............ 11\n'
                     '10. Allowance (allowance) ........................... 12\n'
                     '11. Employee Allowance (employee_allowances) ........ 13\n'
                     '12. Loan (loan) ..................................... 14\n'
                     '13. Loan Payment (loan_payment) ..................... 15\n'
                     '14. Attendance (attendance) ......................... 16\n'
                     '15. Late Computation (late_computation) ............. 17\n'
                     '16. Leave (leave) ................................... 18\n'
                     '17. Leave Type (leave_type) ......................... 19\n'
                     '18. Leave Credit (leave_credit) ..................... 20\n'
                     '19. Leave Credit History (leave_credit_history) ..... 21\n'
                     '20. Department (department) ......................... 22\n'
                     '21. Position (position) ............................. 23\n'
                     '22. Employment Type (employment_type) ............... 24\n'
                     '23. Shift (shift) ................................... 25\n'
                     '24. Employee Shift (employee_shift) ................. 26\n'
                     '25. Job History (job_history) ....................... 27\n'
                     '26. Tax (tax) ....................................... 28')
    
    doc.add_page_break()
    
    # ===== HELPER FUNCTIONS =====
    
    def add_page_break_before_table(paragraph):
        """Add page break property to ensure table starts on new page"""
        p = paragraph._element
        pb = OxmlElement('w:pageBreakBefore')
        pb.set(qn('w:val'), 'true')
        p.insert(0, pb)
    
    def prevent_row_breaks(table):
        """Prevent table rows from splitting across pages"""
        for row in table.rows:
            tr = row._element
            # Keep row together on one page
            keep = OxmlElement('w:keepNext')
            tr.append(keep)
            keep_next = OxmlElement('w:keepLines')
            tr.append(keep_next)
    
    def add_model_table(doc, model_name, table_name, description, fields, relationships=None, notes=None):
        """Add a complete model documentation section on its own page"""
        
        # Force new page for each model
        doc.add_page_break()
        
        # Model Header
        heading = doc.add_heading(model_name, level=1)
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].bold = True
        
        # Table name badge
        table_badge = doc.add_paragraph(f'Table: {table_name}')
        table_badge.runs[0].font.size = Pt(10)
        table_badge.runs[0].italic = True
        table_badge.style = 'Intense Quote'
        
        # Detailed Description
        desc_para = doc.add_paragraph(description)
        desc_para.runs[0].font.size = Pt(10)
        
        # Relationships section if provided
        if relationships:
            rel_heading = doc.add_paragraph('🔗 Relationships:')
            rel_heading.runs[0].bold = True
            rel_heading.runs[0].font.size = Pt(9)
            for rel in relationships:
                rel_para = doc.add_paragraph(f'  • {rel}', style='List Bullet')
                rel_para.runs[0].font.size = Pt(9)
        
        # Additional Notes if provided
        if notes:
            notes_heading = doc.add_paragraph('📝 Notes:')
            notes_heading.runs[0].bold = True
            notes_heading.runs[0].font.size = Pt(9)
            notes_para = doc.add_paragraph(notes)
            notes_para.runs[0].font.size = Pt(9)
            notes_para.style = 'Intense Quote'
        
        doc.add_paragraph()  # Spacing
        
        # Create Table with 5 columns
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set column widths for better fit
        table.columns[0].width = Inches(1.3)   # Column name
        table.columns[1].width = Inches(0.9)   # Type
        table.columns[2].width = Inches(0.5)   # Null
        table.columns[3].width = Inches(0.8)   # Default
        table.columns[4].width = Inches(2.5)   # Description
        
        # Header row styling
        hdr_cells = table.rows[0].cells
        headers = ['Column', 'Type', 'Null', 'Default', 'Description']
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows with compact formatting
        for field in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field.get('name', '')
            row_cells[1].text = field.get('type', '')
            row_cells[2].text = field.get('null', '')
            row_cells[3].text = str(field.get('default', ''))
            row_cells[4].text = field.get('description', '')
            
            # Apply compact font size to all cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
        
        # Prevent table from splitting across pages
        prevent_row_breaks(table)
        
        # Add small spacing after table
        doc.add_paragraph()
    
    # ===== MODEL DEFINITIONS =====
    
    # 1. USERS
    user_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key; unique identifier for each system user'},
        {'name': 'email', 'type': 'String(150)', 'null': 'No', 'default': 'NULL', 'description': 'Unique email address used for authentication and password recovery'},
        {'name': 'password', 'type': 'String(150)', 'null': 'No', 'default': 'NULL', 'description': 'Bcrypt-hashed password; never stored in plaintext'},
        {'name': 'first_name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': "User's legal first name for display and reports"},
        {'name': 'last_name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': "User's legal last name for display and reports"},
        {'name': 'role', 'type': 'String(50)', 'null': 'No', 'default': 'employee', 'description': 'Access role: admin, staff, employee, officer, dept_head - controls UI and permissions'},
        {'name': 'department_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'Foreign key to Department; user\'s assigned department for reporting hierarchy'},
        {'name': 'position', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Job title or position name for organizational charts'},
        {'name': 'active', 'type': 'Boolean', 'null': 'Yes', 'default': 'True', 'description': 'Account status; False disables login without deleting historical data'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Timestamp of account creation for audit trails'},
        {'name': 'last_login', 'type': 'DateTime', 'null': 'Yes', 'default': 'NULL', 'description': 'Timestamp of most recent successful authentication'},
        {'name': 'otp_code', 'type': 'String(10)', 'null': 'Yes', 'default': 'NULL', 'description': '6-digit one-time password for two-factor authentication'},
        {'name': 'otp_expiry', 'type': 'DateTime', 'null': 'Yes', 'default': 'NULL', 'description': 'Expiration timestamp for OTP; typically 5-10 minutes from generation'},
        {'name': 'otp_verified', 'type': 'Boolean', 'null': 'Yes', 'default': 'False', 'description': 'Flag indicating successful OTP verification for current session'},
    ]
    add_model_table(doc, '1. Users', 'users', 
                   'Core authentication table storing all system accounts. Each record represents an individual who can log into GOVHRPAY. The role-based access control (RBAC) system uses the role field to determine dashboard layout, available features, and data visibility. Passwords are hashed using bcrypt before storage. Two-factor authentication (2FA) is supported via OTP fields.',
                   user_fields,
                   relationships=['One-to-One with Employee (employee_profile)', 'One-to-Many with Department (as head)', 'One-to-Many with Leave (as approver)'],
                   notes='• Email must be unique across all users\n• Password reset tokens stored separately\n• Soft delete via active=False, not hard delete\n• Audit logging recommended for role changes')
    
    # 2. EMPLOYEES
    employee_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key; internal system identifier'},
        {'name': 'employee_id', 'type': 'String(20)', 'null': 'No', 'default': 'NULL', 'description': 'Company-assigned employee number; unique business identifier displayed on payslips'},
        {'name': 'user_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to User; links employee record to system login account (nullable for non-system users)'},
        {'name': 'department_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Department; primary organizational assignment for reporting and payroll grouping'},
        {'name': 'position_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Position; current job title for organizational structure and salary banding'},
        {'name': 'first_name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': 'Legal first name as per government ID; used in official documents'},
        {'name': 'last_name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': 'Legal last name as per government ID; used in official documents'},
        {'name': 'middle_name', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Middle name or initial; required for Philippine government compliance'},
        {'name': 'email', 'type': 'String(150)', 'null': 'No', 'default': 'NULL', 'description': 'Corporate email address; must be unique; used for payslip distribution'},
        {'name': 'phone', 'type': 'String(20)', 'null': 'Yes', 'default': 'NULL', 'description': 'Primary contact number for SMS notifications and emergency contact'},
        {'name': 'barangay', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Barangay (village) of residence; required for Philippine address format'},
        {'name': 'municipality', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'City or municipality of residence for tax jurisdiction determination'},
        {'name': 'province', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Province of residence for regional reporting and tax calculations'},
        {'name': 'postal_code', 'type': 'String(10)', 'null': 'Yes', 'default': 'NULL', 'description': 'Philippine postal/ZIP code for mailing and geographic grouping'},
        {'name': 'street_address', 'type': 'String(255)', 'null': 'Yes', 'default': 'NULL', 'description': 'House number, street, subdivision; complete mailing address'},
        {'name': 'salary', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Current basic monthly salary; base for all payroll calculations and deductions'},
        {'name': 'date_hired', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'Official start date; used for tenure calculation, leave accrual, and probation tracking'},
        {'name': 'date_of_birth', 'type': 'Date', 'null': 'Yes', 'default': 'NULL', 'description': 'Birth date for age verification, retirement planning, and benefits eligibility'},
        {'name': 'gender', 'type': 'String(10)', 'null': 'Yes', 'default': 'NULL', 'description': 'Gender identity for demographic reporting and maternity/paternity leave rules'},
        {'name': 'marital_status', 'type': 'String(20)', 'null': 'Yes', 'default': 'NULL', 'description': 'Single/Married/etc.; affects tax exemptions and dependent benefits'},
        {'name': 'emergency_contact', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Name of emergency contact person for workplace safety protocols'},
        {'name': 'emergency_phone', 'type': 'String(20)', 'null': 'Yes', 'default': 'NULL', 'description': 'Emergency contact phone number for urgent communications'},
        {'name': 'status', 'type': 'String(20)', 'null': 'Yes', 'default': 'Active', 'description': 'Employment status: Active, On-Leave, Resigned, Terminated; controls payroll inclusion'},
        {'name': 'archived', 'type': 'Boolean', 'null': 'Yes', 'default': 'False', 'description': 'Soft-delete flag; archived employees excluded from active lists but retained for history'},
        {'name': 'archived_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'NULL', 'description': 'Timestamp when employee was archived for audit and rehire tracking'},
        {'name': 'cs_eligibility', 'type': 'String(50)', 'null': 'Yes', 'default': 'NULL', 'description': 'Civil Service eligibility status; required for government-mandated positions'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Record creation timestamp for data governance'},
        {'name': 'updated_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Last modification timestamp for change tracking'},
        {'name': 'employment_type_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to EmploymentType; determines payroll rules (Regular, Part-Time, Casual, Job Order)'},
    ]
    add_model_table(doc, '2. Employees', 'employee',
                   'Master employee record containing comprehensive personal, contact, and employment data. This is the central entity for HR operations. Each employee may optionally have a linked User account for system access. The employment_type_id field drives payroll calculation logic (e.g., daily rate vs hourly rate). Address fields follow Philippine format for compliance. Status field controls inclusion in active payroll runs.',
                   employee_fields,
                   relationships=['One-to-One with User (user_id)', 'Many-to-One with Department', 'Many-to-One with Position', 'Many-to-One with EmploymentType', 'One-to-Many with Attendance, Leave, Payroll, Loan'],
                   notes='• employee_id is business key; id is technical key\n• salary changes tracked via JobHistory, not updated in-place\n• archived employees retain historical payroll data\n• email unique constraint applies only to non-null values')
    
    # 3. PAYROLL PERIOD
    pp_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for payroll period'},
        {'name': 'period_name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': 'Human-readable label (e.g., "April 1-15, 2026"); displayed on payslips and reports'},
        {'name': 'start_date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'First day of the pay period; used to filter attendance and leave records'},
        {'name': 'end_date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'Last day of the pay period; inclusive endpoint for date-range queries'},
        {'name': 'pay_date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'Scheduled payment/disbursement date; used for bank file generation and cash flow forecasting'},
        {'name': 'status', 'type': 'String(30)', 'null': 'Yes', 'default': 'Open', 'description': 'Workflow state: Open (editable), Processing (locked), Closed (finalized), Paid (disbursed)'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Period creation timestamp for audit'},
    ]
    add_model_table(doc, '3. Payroll Period', 'payroll_period',
                   'Defines discrete payroll cycles (semi-monthly, monthly, etc.). Each period groups employee payroll records for batch processing. The status field implements a state machine: Open periods allow payroll edits; Processing locks calculations; Closed prevents changes; Paid marks disbursement completion. Date ranges are used to aggregate attendance, overtime, and leave data for pay computation.',
                   pp_fields,
                   relationships=['One-to-Many with Payroll (period.payrolls)', 'Used by Payslip generation'],
                   notes='• Periods should not overlap\n• end_date must be >= start_date\n• pay_date typically 3-5 days after end_date\n• Status transitions should be logged')
    
    # 4. PAYROLL
    payroll_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for individual payroll record'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; identifies the employee being paid'},
        {'name': 'payroll_period_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to PayrollPeriod; groups this record into a pay cycle'},
        {'name': 'basic_salary', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Prorated basic salary for the period; computed from employee.salary based on days/hours worked'},
        {'name': 'working_hours', 'type': 'Float', 'null': 'Yes', 'default': '160', 'description': 'Standard hours for the period (e.g., 20 days × 8 hrs); basis for hourly rate calculation'},
        {'name': 'hours_worked', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Actual hours from Attendance records; used for Part-Time/Casual pay computation'},
        {'name': 'days_worked', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Count of Present/Late days from Attendance; used for Regular/Casual daily rate pay'},
        {'name': 'overtime_hours', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Approved overtime hours; multiplied by 1.25× hourly rate per labor code'},
        {'name': 'holiday_pay', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Holiday premium pay for work on regular/special holidays'},
        {'name': 'night_diff', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Night differential (10% premium) for hours worked between 10PM-6AM'},
        {'name': 'gross_pay', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Sum of basic_salary + overtime_pay + holiday_pay + night_diff + allowances'},
        {'name': 'total_deductions', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Aggregate of all employee_share deductions from PayrollDeduction records'},
        {'name': 'net_pay', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Final take-home pay: gross_pay - total_deductions; amount disbursed to employee'},
        {'name': 'status', 'type': 'String(30)', 'null': 'Yes', 'default': 'Draft', 'description': 'Record state: Draft (editable), Approved (locked), Paid (disbursed)'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Payroll record creation timestamp'},
        {'name': 'allowance_total', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Sum of all applicable allowances from EmployeeAllowance records'},
    ]
    add_model_table(doc, '4. Payroll', 'payroll',
                   'Individual employee payroll calculation record. This is the core transactional table for pay computation. The calculate() method aggregates attendance data, applies employment-type-specific formulas, computes earnings (basic, OT, holiday, night diff), adds allowances, then subtracts deductions to produce net_pay. Status workflow: Draft → Approved (locks values) → Paid (triggers disbursement). Each record links to detailed PayrollDeduction breakdown.',
                   payroll_fields,
                   relationships=['Many-to-One with Employee', 'Many-to-One with PayrollPeriod', 'One-to-One with Payslip', 'One-to-Many with PayrollDeduction', 'One-to-Many with LoanPayment'],
                   notes='• gross_pay, total_deductions, net_pay are computed, not user-input\n• Recalculation requires status=Draft\n• Historical records should never be modified after Paid\n• Use database transactions for atomic calculation+save')
    
    # 5. PAYROLL DEDUCTION
    pd_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for deduction line item'},
        {'name': 'payroll_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Payroll; parent payroll record this deduction belongs to'},
        {'name': 'deduction_name', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Display name of deduction (e.g., "SSS", "PhilHealth", "Pag-IBIG", "Tax")'},
        {'name': 'employee_share', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Amount deducted from employee\'s net pay; contributes to total_deductions'},
        {'name': 'employer_share', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Company contribution portion; tracked for statutory reporting but not deducted from employee'},
        {'name': 'ec', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Employee Compensation (EC) fund contribution; specific to Philippine SSS'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Deduction record creation timestamp'},
    ]
    add_model_table(doc, '5. Payroll Deduction', 'payroll_deduction',
                   'Itemized deduction breakdown for each payroll record. Created automatically during Payroll.calculate() by iterating through employee\'s active EmployeeDeduction records and applying calculation rules. Stores both employee_share (deducted from pay) and employer_share (company cost) for complete statutory reporting. The ec field handles Philippine-specific SSS Employee Compensation contributions.',
                   pd_fields,
                   relationships=['Many-to-One with Payroll (backref: deduction_breakdown)', 'Derived from EmployeeDeduction + Deduction rules'],
                   notes='• Records are cascade-deleted when parent Payroll is deleted\n• employee_share values summed to form Payroll.total_deductions\n• For audit: store calculation inputs/outputs separately if needed')
    
    # 6. PAYSLIP
    payslip_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for payslip document'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; recipient of the payslip'},
        {'name': 'payroll_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Payroll; source of pay data; one-to-one relationship'},
        {'name': 'payslip_number', 'type': 'String(50)', 'null': 'Yes', 'default': 'NULL', 'description': 'Unique reference number (e.g., "PS-2026-04-001"); used for employee inquiries and audit'},
        {'name': 'gross_pay', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Snapshot of payroll gross_pay at time of generation'},
        {'name': 'total_deductions', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Snapshot of payroll total_deductions at time of generation'},
        {'name': 'net_pay', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Snapshot of payroll net_pay at time of generation'},
        {'name': 'generated_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Timestamp when payslip PDF/document was created'},
        {'name': 'status', 'type': 'String(30)', 'null': 'Yes', 'default': 'Generated', 'description': 'Document status: Generated, Sent, Viewed, Downloaded'},
    ]
    add_model_table(doc, '6. Payslip', 'payslip',
                   'Immutable snapshot of payroll data for employee distribution. Created when payroll is approved; captures gross_pay, deductions, and net_pay at that moment to preserve historical accuracy even if source payroll is later corrected. The payslip_number provides a human-readable reference. Status tracking enables delivery confirmation (email sent, employee viewed). One-to-one with Payroll ensures each payroll has exactly one official payslip.',
                   payslip_fields,
                   relationships=['Many-to-One with Employee', 'One-to-One with Payroll (uselist=False)'],
                   notes='• Payslip values should never change after generation\n• Use separate PDF generation service for actual document creation\n• Consider encryption for stored payslip files\n• Status updates should trigger notifications')
    
    # 7. DEDUCTION
    ded_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for deduction type definition'},
        {'name': 'name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': 'Unique deduction name (e.g., "SSS", "PhilHealth", "Pag-IBIG", "Withholding Tax")'},
        {'name': 'description', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Detailed explanation of deduction purpose, legal basis, and calculation methodology'},
        {'name': 'calculation_type', 'type': 'String(20)', 'null': 'No', 'default': 'NULL', 'description': 'Formula type: fixed (flat amount), percentage (% of salary), bracket (salary-range table), progressive (tiered %)'},
        {'name': 'rate', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Percentage rate for percentage/progressive types (e.g., 0.045 for 4.5% SSS)'},
        {'name': 'ceiling', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Maximum salary base for calculation (e.g., SSS contribution capped at ₱30,000)'},
        {'name': 'floor', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Minimum salary threshold before deduction applies'},
        {'name': 'active', 'type': 'Boolean', 'null': 'Yes', 'default': 'True', 'description': 'Enable/disable deduction globally without deleting historical references'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Deduction type creation timestamp'},
    ]
    add_model_table(doc, '7. Deduction', 'deduction',
                   'Master definition of statutory and voluntary deductions. Supports four calculation strategies: (1) fixed: flat amount regardless of salary; (2) percentage: rate × salary with optional floor/ceiling; (3) bracket: lookup table by salary range; (4) progressive: tiered percentage with fixed base amount. Used by EmployeeDeduction.calculate() during payroll processing. active flag allows temporary suspension (e.g., loan completion) without losing configuration.',
                   ded_fields,
                   relationships=['One-to-Many with DeductionBracket (for bracket/progressive types)', 'One-to-Many with EmployeeDeduction (employee assignments)'],
                   notes='• name must be unique; use codes for internal references\n• Changing calculation_type requires careful migration of existing EmployeeDeductions\n• Ceiling/floor applied before rate multiplication\n• Progressive type uses DeductionBracket.rate + fixed_amount')
    
    # 8. EMPLOYEE DEDUCTION
    emp_ded_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for employee-deduction assignment'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; the employee subject to this deduction'},
        {'name': 'deduction_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Deduction; the deduction type/rule to apply'},
        {'name': 'override_amount', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Optional fixed amount that bypasses standard calculation; used for custom arrangements'},
        {'name': 'active', 'type': 'Boolean', 'null': 'Yes', 'default': 'True', 'description': 'Enable/disable this deduction for the employee without deleting the assignment'},
    ]
    add_model_table(doc, '8. Employee Deduction', 'employee_deductions',
                   'Junction table linking employees to deduction types with optional customization. During payroll calculation, active records are processed through Deduction.calculate() unless override_amount is set, in which case the override value is used directly. This allows per-employee exceptions (e.g., salary loan with fixed amortization) while maintaining centralized deduction rules. The active flag supports temporary suspensions (e.g., leave without pay periods).',
                   emp_ded_fields,
                   relationships=['Many-to-One with Employee', 'Many-to-One with Deduction', 'Source for PayrollDeduction records during payroll calculation'],
                   notes='• Unique constraint on (employee_id, deduction_id) recommended\n• override_amount takes precedence over Deduction rules\n• Deactivate instead of delete to preserve history\n• Consider adding start_date/end_date for time-bound deductions')
    
    # 9. DEDUCTION BRACKET
    bracket_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for salary bracket definition'},
        {'name': 'deduction_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Deduction; parent deduction type this bracket belongs to'},
        {'name': 'salary_from', 'type': 'Float', 'null': 'No', 'default': 'NULL', 'description': 'Minimum salary (inclusive) for this bracket range'},
        {'name': 'salary_to', 'type': 'Float', 'null': 'No', 'default': 'NULL', 'description': 'Maximum salary (inclusive) for this bracket range'},
        {'name': 'employee_share', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Fixed deduction amount for bracket-type calculations'},
        {'name': 'employer_share', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Fixed employer contribution for this bracket'},
        {'name': 'ec', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Employee Compensation amount for this bracket (Philippine SSS)'},
        {'name': 'rate', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Percentage rate for progressive-type calculations'},
        {'name': 'fixed_amount', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Base amount added to (salary × rate) for progressive calculations'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Bracket creation timestamp'},
    ]
    add_model_table(doc, '9. Deduction Bracket', 'deduction_bracket',
                   'Salary range definitions for bracket-based and progressive deductions. For bracket type: when employee salary falls within [salary_from, salary_to], the fixed employee_share/employer_share/ec values apply. For progressive type: deduction = (salary × rate) + fixed_amount. Used by Philippine tax tables and SSS contribution schedules where amounts vary by income tier. Multiple brackets per deduction enable complex statutory formulas.',
                   bracket_fields,
                   relationships=['Many-to-One with Deduction (backref: brackets)'],
                   notes='• Brackets should not overlap; validate salary_from/to ranges\n• Use salary_to = NULL for open-ended top bracket\n• Progressive formula: result = min(salary, salary_to) * rate + fixed_amount\n• Update brackets when government rates change; historical payroll uses rules at time of calculation')
    
    # 10. ALLOWANCE
    allow_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for allowance type'},
        {'name': 'name', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Allowance name (e.g., "Rice Subsidy", "Transportation", "COLA", "Uniform")'},
        {'name': 'amount', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Standard monetary value of the allowance'},
        {'name': 'active', 'type': 'Boolean', 'null': 'Yes', 'default': 'True', 'description': 'Enable/disable allowance globally'},
        {'name': 'min_salary', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Minimum employee salary to qualify for this allowance'},
        {'name': 'max_salary', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Maximum employee salary to qualify (NULL = no upper limit)'},
    ]
    add_model_table(doc, '10. Allowance', 'allowance',
                   'Master definition of employee allowances and benefits. Allowances are added to gross pay during payroll calculation. Eligibility can be restricted by salary range using min_salary/max_salary fields (e.g., transportation allowance only for employees earning below ₱25,000). The active flag allows temporary suspension. Amount field provides default value; EmployeeAllowance can override for individual exceptions.',
                   allow_fields,
                   relationships=['One-to-Many with EmployeeAllowance (employee assignments)'],
                   notes='• Allowances are non-taxable up to legal limits; track separately for BIR reporting\n• Consider adding frequency field (monthly, per-day, one-time)\n• Salary eligibility checked at payroll time, not assignment time')
    
    # 11. EMPLOYEE ALLOWANCE
    emp_allow_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for employee-allowance assignment'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Employee; recipient of the allowance'},
        {'name': 'allowance_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Allowance; type of allowance granted'},
    ]
    add_model_table(doc, '11. Employee Allowance', 'employee_allowances',
                   'Junction table assigning allowances to employees. During payroll calculation, active assignments are summed into allowance_total. Simple many-to-many relationship; complex eligibility (salary ranges, department, employment type) is evaluated at calculation time using the parent Allowance rules. For custom amounts per employee, consider adding an override_amount field similar to EmployeeDeduction.',
                   emp_allow_fields,
                   relationships=['Many-to-One with Employee', 'Many-to-One with Allowance'],
                   notes='• Unique constraint on (employee_id, allowance_id) recommended\n• Consider adding effective_date/expiration_date for time-bound allowances\n• Deactivate assignments instead of deleting to preserve history')
    
    # 12. LOAN
    loan_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for employee loan record'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Employee; borrower of the loan'},
        {'name': 'provider', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Lending institution: Pag-IBIG, SSS, Bank, Company, etc.'},
        {'name': 'loan_type', 'type': 'String(100)', 'null': 'Yes', 'default': 'NULL', 'description': 'Loan product: Multi-Purpose Loan (MPL), Salary Loan, Emergency Loan, Housing Loan'},
        {'name': 'total_amount', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Original principal amount approved'},
        {'name': 'monthly_payment', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Scheduled amortization amount deducted per payroll period'},
        {'name': 'remaining_balance', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Current outstanding principal; updated after each LoanPayment'},
        {'name': 'start_date', 'type': 'Date', 'null': 'Yes', 'default': 'NULL', 'description': 'Date when amortization deductions begin'},
        {'name': 'active', 'type': 'Boolean', 'null': 'Yes', 'default': 'True', 'description': 'Loan status: True = active deductions; False = paid-off or suspended'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Loan application/approval timestamp'},
    ]
    add_model_table(doc, '12. Loan', 'loan',
                   'Employee loan records for salary-deducted amortizations. Supports government loans (Pag-IBIG, SSS) and company loans. During payroll calculation, active loans with start_date <= period end_date generate LoanPayment records with amount = monthly_payment. remaining_balance is decremented after each payment; when balance <= 0, active is set to False. Provider and loan_type fields enable reporting by lender and product.',
                   loan_fields,
                   relationships=['Many-to-One with Employee', 'One-to-Many with LoanPayment'],
                   notes='• Validate monthly_payment × terms ≈ total_amount during creation\n• remaining_balance should never be negative\n• Consider adding interest_rate, term_months, penalty_rate for complex loans\n• Archive paid loans instead of deleting for audit')
    
    # 13. LOAN PAYMENT
    loan_pay_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for loan payment transaction'},
        {'name': 'loan_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Loan; the loan being paid'},
        {'name': 'payroll_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Payroll; the payroll period when deduction occurred'},
        {'name': 'amount_paid', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Actual amount deducted from employee\'s net pay this period'},
        {'name': 'remaining_balance', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Loan balance after this payment; snapshot for historical accuracy'},
        {'name': 'payment_date', 'type': 'Date', 'null': 'Yes', 'default': 'NULL', 'description': 'Date payment was processed (typically payroll pay_date)'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Payment record creation timestamp'},
    ]
    add_model_table(doc, '13. Loan Payment', 'loan_payment',
                   'Individual loan amortization transactions linked to payroll. Created automatically during Payroll.calculate() for each active Loan meeting eligibility criteria. amount_paid is deducted from net_pay and added to total_deductions. remaining_balance provides an immutable snapshot of the loan state at payment time, enabling accurate historical reporting even if the parent Loan record is later modified. payment_date aligns with payroll disbursement for bank reconciliation.',
                   loan_pay_fields,
                   relationships=['Many-to-One with Loan', 'Many-to-One with Payroll'],
                   notes='• amount_paid should match Loan.monthly_payment unless partial payment allowed\n• remaining_balance = previous_balance - amount_paid\n• Use database transaction to update Loan.remaining_balance atomically\n• Consider adding interest_paid, principal_paid split for amortization schedules')
    
    # 14. ATTENDANCE
    att_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for attendance record'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; the employee clocking in/out'},
        {'name': 'date', 'type': 'Date', 'null': 'No', 'default': 'today', 'description': 'Calendar date of the attendance; unique per employee (enforced by application logic)'},
        {'name': 'time_in', 'type': 'Time', 'null': 'Yes', 'default': 'NULL', 'description': 'Actual clock-in timestamp; used for late calculation and working hours'},
        {'name': 'time_out', 'type': 'Time', 'null': 'Yes', 'default': 'NULL', 'description': 'Actual clock-out timestamp; used for working hours calculation'},
        {'name': 'status', 'type': 'String(50)', 'null': 'Yes', 'default': 'Present', 'description': 'Attendance status: Present, Late, Absent, Leave, Holiday, Rest Day'},
        {'name': 'remarks', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Manual notes: tardiness reason, approved overtime, system override justification'},
        {'name': 'working_hours', 'type': 'Float', 'null': 'Yes', 'default': '0.0', 'description': 'Calculated productive hours: min(time_out, shift_end) - max(time_in, shift_start), minus 1hr break if >4hrs'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Record creation timestamp (biometric sync or manual entry time)'},
    ]
    add_model_table(doc, '14. Attendance', 'attendance',
                   'Daily time records capturing employee presence and work duration. Created via biometric device sync, mobile app, or manual entry. The calculate_working_hours() method computes productive time based on assigned shift, excluding meal breaks. Status is auto-set by check_late(): if time_in > shift.start_time, status=Late. working_hours feeds into Part-Time/Casual payroll calculations. Remarks field supports audit trails for manual adjustments.',
                   att_fields,
                   relationships=['Many-to-One with Employee', 'One-to-One with LateComputation (auto-generated for late arrivals)'],
                   notes='• Unique constraint on (employee_id, date) recommended\n• working_hours auto-calculated via SQLAlchemy event listeners\n• Late status triggers LateComputation record creation\n• Absent records may be auto-generated for missing days in pay period')
    
    # 15. LATE COMPUTATION
    late_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for late calculation record'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; the employee who arrived late'},
        {'name': 'attendance_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Attendance; source record with time_in > scheduled start'},
        {'name': 'date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'Date of late arrival; denormalized for easier querying'},
        {'name': 'late_days', 'type': 'Integer', 'null': 'Yes', 'default': '0', 'description': 'Number of full days late (rare; typically 0)'},
        {'name': 'late_hours', 'type': 'Integer', 'null': 'Yes', 'default': '0', 'description': 'Hours portion of tardiness (e.g., 2 for 2 hours late)'},
        {'name': 'late_minutes', 'type': 'Integer', 'null': 'Yes', 'default': '0', 'description': 'Minutes portion of tardiness (e.g., 30 for 30 minutes late)'},
        {'name': 'day_equivalent', 'type': 'Float', 'null': 'No', 'default': 'NULL', 'description': 'Converted to work days: days×1.0 + hours×0.125 + minutes×0.002; used for salary deduction'},
        {'name': 'remarks', 'type': 'String(255)', 'null': 'Yes', 'default': 'NULL', 'description': 'Notes on late reason, approval for excused tardiness, or calculation override'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Record creation timestamp'},
    ]
    add_model_table(doc, '15. Late Computation', 'late_computation',
                   'Detailed tardiness calculations converting time-based lateness into day equivalents for payroll deduction. Auto-generated by event listener when Attendance.time_in > scheduled shift start. Uses Philippine labor formula: 1 hour = 0.125 days (1/8 of workday), 1 minute = 0.002 days (Excel-compatible rounding). The day_equivalent value is subtracted from days_worked during payroll calculation for Regular/Casual employees. Remarks support excused late approvals.',
                   late_fields,
                   relationships=['Many-to-One with Employee', 'One-to-One with Attendance'],
                   notes='• day_equivalent formula: (days*1.0) + (hours*0.125) + (minutes*0.002)\n• Rounded to 3 decimals for precision\n• Only created for Late status attendances\n• Consider adding approved_by field for manager-excused tardiness')
    
    # 16. LEAVE
    leave_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for leave request'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; the employee requesting leave'},
        {'name': 'leave_type_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to LeaveType; category of leave (Vacation, Sick, Maternity, etc.)'},
        {'name': 'start_date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'First day of leave; inclusive'},
        {'name': 'end_date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'Last day of leave; inclusive'},
        {'name': 'days_requested', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'Total calendar days requested (end_date - start_date + 1)'},
        {'name': 'reason', 'type': 'Text', 'null': 'No', 'default': 'NULL', 'description': 'Employee-provided justification for leave request'},
        {'name': 'status', 'type': 'String(50)', 'null': 'Yes', 'default': 'Pending', 'description': 'Workflow state: Pending, Approved, Denied, Cancelled, Withdrawn'},
        {'name': 'approved_by', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to User; manager who approved/denied the request'},
        {'name': 'approved_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'NULL', 'description': 'Timestamp of approval/denial decision'},
        {'name': 'comments', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Approver feedback or denial reason'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Request submission timestamp'},
        {'name': 'updated_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Last status change timestamp'},
        {'name': 'paid_days', 'type': 'Integer', 'null': 'Yes', 'default': '0', 'description': 'Number of days covered by paid leave credits'},
        {'name': 'unpaid_days', 'type': 'Integer', 'null': 'Yes', 'default': '0', 'description': 'Number of days exceeding paid credits; deducted from salary'},
        {'name': 'canceled_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'NULL', 'description': 'Timestamp if employee withdrew request before approval'},
    ]
    add_model_table(doc, '16. Leave', 'leave',
                   'Employee leave requests with approval workflow. Supports multiple leave types with different rules (max paid days, documentation requirements). Upon approval, compute_paid_leave() splits days into paid_days (deducted from LeaveCredit) and unpaid_days (salary deduction). Status workflow: Pending → Approved/Denied → (optional) Cancelled. Approved leaves reduce available credits and affect payroll days_worked calculation. comments field captures approver rationale for audit.',
                   leave_fields,
                   relationships=['Many-to-One with Employee', 'Many-to-One with LeaveType', 'Many-to-One with User (approver)'],
                   notes='• Validate days_requested = (end_date - start_date + 1) excluding weekends/holidays if policy requires\n• paid_days + unpaid_days should equal days_requested\n• Approved leaves should lock credits to prevent double-booking\n• Consider adding attachment field for medical certificates')
    
    # 17. LEAVE TYPE
    lt_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for leave category'},
        {'name': 'name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': 'Unique leave type name: Vacation, Sick, Maternity, Paternity, Bereavement, Emergency'},
        {'name': 'description', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Detailed policy: eligibility, documentation required, notice period, carryover rules'},
        {'name': 'max_paid_days', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'Maximum days payable from leave credits per year (e.g., 15 for Vacation, 105 for Maternity)'},
        {'name': 'max_duration_days', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'Maximum consecutive days allowed per request (e.g., 30 for Vacation, 105 for Maternity)'},
    ]
    add_model_table(doc, '17. Leave Type', 'leave_type',
                   'Master definitions of leave categories with policy rules. Each type defines: (1) max_paid_days: annual entitlement deducted from LeaveCredit; (2) max_duration_days: longest single request allowed. Used during leave request validation and payroll calculation to determine paid vs unpaid split. Descriptions should reference company policy documents and legal requirements (e.g., Philippine Labor Code maternity provisions).',
                   lt_fields,
                   relationships=['One-to-Many with Leave', 'One-to-Many with LeaveCredit'],
                   notes='• name should be unique and human-readable\n• max_paid_days=NULL means unlimited paid leave (rare)\n• Consider adding accrual_rate field for automatic credit generation\n• Policy changes should not affect historical leave records')
    
    # 18. LEAVE CREDIT
    lc_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for employee leave balance'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; the employee holding the credits'},
        {'name': 'leave_type_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to LeaveType; the leave category these credits apply to'},
        {'name': 'total_credits', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Cumulative earned credits for this leave type (increases via accrual or manual grant)'},
        {'name': 'used_credits', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Cumulative consumed credits (increases when approved leave is taken)'},
    ]
    add_model_table(doc, '18. Leave Credit', 'leave_credit',
                   'Employee-specific leave balances per leave type. total_credits accumulates via annual accrual, promotions, or manual adjustments. used_credits increments when Approved leaves are taken. remaining_credits() = total - used determines eligibility for new requests. Updated transactionally when leaves are approved to prevent overuse. One record per employee per leave type; created automatically on hire or leave type activation.',
                   lc_fields,
                   relationships=['Many-to-One with Employee', 'Many-to-One with LeaveType', 'Source for LeaveCreditHistory entries'],
                   notes='• Unique constraint on (employee_id, leave_type_id)\n• used_credits should never exceed total_credits (enforce in application)\n• Consider adding fiscal_year field for annual reset policies\n• Audit changes via LeaveCreditHistory, not by updating this record')
    
    # 19. LEAVE CREDIT HISTORY
    lch_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for credit transaction'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; affected employee'},
        {'name': 'leave_type_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to LeaveType; affected leave category'},
        {'name': 'earned', 'type': 'Float', 'null': 'Yes', 'default': '0.0', 'description': 'Credits added this period (positive value); from accrual, promotion, or manual grant'},
        {'name': 'used', 'type': 'Float', 'null': 'Yes', 'default': '0.0', 'description': 'Credits consumed this period (positive value); from approved leave usage'},
        {'name': 'month', 'type': 'String(20)', 'null': 'Yes', 'default': 'NULL', 'description': 'Period label (e.g., "Apr 2026") for grouping and reporting'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Transaction timestamp'},
    ]
    add_model_table(doc, '19. Leave Credit History', 'leave_credit_history',
                   'Immutable audit trail of leave credit changes. Each record captures a single transaction: either earned (credit grant) or used (leave consumption). Enables balance reconciliation, dispute resolution, and trend analysis. Created automatically when: (1) annual accrual runs, (2) leave is approved, (3) admin manually adjusts credits. The month field supports monthly reporting without complex date calculations.',
                   lch_fields,
                   relationships=['Many-to-One with Employee', 'Many-to-One with LeaveType'],
                   notes='• earned and used are mutually exclusive per record (one is 0)\n• Net change = earned - used; cumulative sum should match LeaveCredit\n• Consider adding reference_id to link to source (Leave.id, batch job ID)\n• Never delete history records; use soft-delete if absolutely necessary')
    
    # 20. DEPARTMENT
    dept_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for organizational unit'},
        {'name': 'name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': 'Unique department name: HR, Finance, IT, Operations, etc.'},
        {'name': 'description', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Department mandate, key responsibilities, and reporting structure'},
        {'name': 'head_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to User; department manager who approves leaves and views reports'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Department creation timestamp'},
    ]
    add_model_table(doc, '20. Department', 'department',
                   'Organizational units for grouping employees and delegating approval authority. Used for: (1) filtering employee lists and reports, (2) routing leave approvals to department heads, (3) cost center allocation in payroll. The head_id field identifies the manager with approval permissions. Departments can be hierarchical in application logic (parent_id field) if needed, though not enforced in schema.',
                   dept_fields,
                   relationships=['One-to-Many with Employee', 'One-to-Many with Position', 'One-to-Many with User (as head)'],
                   notes='• name must be unique; use codes for internal references if needed\n• Changing head_id should trigger permission updates\n• Consider adding cost_center_code for financial integration\n• Archive instead of delete to preserve historical employee assignments')
    
    # 21. POSITION
    pos_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for job position'},
        {'name': 'name', 'type': 'String(100)', 'null': 'No', 'default': 'NULL', 'description': 'Unique position title: Software Engineer, HR Officer, Accountant, etc.'},
        {'name': 'description', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Key responsibilities, required qualifications, and reporting relationships'},
        {'name': 'department_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Department; organizational unit this position belongs to'},
    ]
    add_model_table(doc, '21. Position', 'position',
                   'Job position definitions for organizational structure and salary banding. Used to: (1) categorize employees for reporting, (2) enforce salary ranges during hiring/promotions, (3) define approval hierarchies. Each position belongs to one department; employees reference positions for their current role. Position changes are tracked via JobHistory, not by updating Employee.position_id directly.',
                   pos_fields,
                   relationships=['Many-to-One with Department', 'One-to-Many with Employee'],
                   notes='• name should be unique; consider adding level/seniority field\n• Description should align with official job descriptions\n• Position changes should create JobHistory records, not modify Employee directly\n• Consider adding salary_min/salary_max for compensation governance')
    
    # 22. EMPLOYMENT TYPE
    et_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for employment classification'},
        {'name': 'name', 'type': 'String(50)', 'null': 'No', 'default': 'NULL', 'description': 'Classification: Regular, Part-Time, Casual, Job Order (JO), Contractual'},
        {'name': 'description', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Legal definition, benefits eligibility, and payroll calculation rules'},
    ]
    add_model_table(doc, '22. Employment Type', 'employment_type',
                   'Employment classifications that drive payroll calculation logic and benefits eligibility. Each type defines: (1) pay basis (monthly salary vs hourly vs daily), (2) leave accrual rules, (3) statutory benefit requirements. Referenced by Employee.employment_type_id to determine which formula Payroll.calculate() applies. Critical for Philippine labor compliance where Regular, Casual, and Job Order workers have different legal protections.',
                   et_fields,
                   relationships=['One-to-Many with Employee'],
                   notes='• name values must match strings used in Payroll.calculate() conditionals\n• Consider adding is_monthly, is_hourly boolean flags for clearer logic\n• Changing an employee\'s type should be tracked in JobHistory\n• Some types may be phased out; use active flag instead of deletion')
    
    # 23. SHIFT
    shift_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for work shift definition'},
        {'name': 'name', 'type': 'String(50)', 'null': 'No', 'default': 'NULL', 'description': 'Shift label: Morning, Afternoon, Night, Rotating'},
        {'name': 'start_time', 'type': 'Time', 'null': 'No', 'default': 'NULL', 'description': 'Scheduled start time (e.g., 08:00:00 for morning shift)'},
        {'name': 'end_time', 'type': 'Time', 'null': 'No', 'default': 'NULL', 'description': 'Scheduled end time (e.g., 17:00:00 for 8-hour shift with 1hr break)'},
    ]
    add_model_table(doc, '23. Shift', 'shift',
                   'Work schedule templates defining standard start/end times. Used by Attendance.check_late() to determine tardiness and by Attendance.calculate_working_hours() to compute productive time. Shifts can be assigned to employees via EmployeeShift for daily variations or to positions for default schedules. Time fields store time-of-day only; date context comes from Attendance or EmployeeShift records.',
                   shift_fields,
                   relationships=['One-to-Many with EmployeeShift'],
                   notes='• end_time may be < start_time for overnight shifts (handle in application logic)\n• Consider adding break_duration field for automatic working_hours adjustment\n• Shift changes should be tracked via EmployeeShift, not by modifying Shift\n• Validate that end_time - start_time matches expected work hours')
    
    # 24. EMPLOYEE SHIFT
    es_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for daily shift assignment'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; the employee assigned to this shift'},
        {'name': 'shift_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Shift; the schedule template to apply'},
        {'name': 'date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'Calendar date this shift assignment applies to'},
        {'name': 'day_of_week', 'type': 'String(15)', 'null': 'No', 'default': 'NULL', 'description': 'Day name (Monday, Tuesday, etc.) for quick filtering and reporting'},
        {'name': 'status', 'type': 'String(15)', 'null': 'No', 'default': 'active', 'description': 'Assignment status: active, cancelled, swapped (for schedule change tracking)'},
    ]
    add_model_table(doc, '24. Employee Shift', 'employee_shift',
                   'Daily shift assignments overriding an employee\'s default schedule. Enables flexible scheduling: rotating shifts, holiday coverage, temporary reassignments. Referenced by Attendance.get_shift() to determine the scheduled start_time for late calculation. The unique (employee_id, date) constraint ensures one shift per day. status field supports schedule change workflows without deleting historical assignments.',
                   es_fields,
                   relationships=['Many-to-One with Employee', 'Many-to-One with Shift'],
                   notes='• Unique constraint on (employee_id, date) enforced via __table_args__\n• day_of_week is denormalized for easier querying; validate against date\n• Consider adding created_by field for audit of schedule changes\n• Cancelled shifts should not generate Attendance records')
    
    # 25. JOB HISTORY
    jh_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for employment history record'},
        {'name': 'employee_id', 'type': 'Integer', 'null': 'No', 'default': 'NULL', 'description': 'FK to Employee; the employee whose history is recorded'},
        {'name': 'effective_date', 'type': 'Date', 'null': 'No', 'default': 'NULL', 'description': 'Start date of this appointment/position/salary change'},
        {'name': 'end_date', 'type': 'Date', 'null': 'Yes', 'default': 'NULL', 'description': 'End date of this record; NULL means current/ongoing'},
        {'name': 'position_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Position; job title during this period'},
        {'name': 'employment_type_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to EmploymentType; classification during this period'},
        {'name': 'department_id', 'type': 'Integer', 'null': 'Yes', 'default': 'NULL', 'description': 'FK to Department; organizational unit during this period'},
        {'name': 'salary', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Basic salary during this period; snapshot for historical payroll accuracy'},
        {'name': 'status', 'type': 'String(50)', 'null': 'Yes', 'default': 'NULL', 'description': 'Employment status during period: Active, Promoted, Transferred, Resigned, Terminated'},
        {'name': 'remarks', 'type': 'Text', 'null': 'Yes', 'default': 'NULL', 'description': 'Reason for change: promotion details, transfer justification, separation cause'},
        {'name': 'created_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'History record creation timestamp'},
        {'name': 'updated_at', 'type': 'DateTime', 'null': 'Yes', 'default': 'utcnow', 'description': 'Last modification timestamp'},
    ]
    add_model_table(doc, '25. Job History', 'job_history',
                   'Immutable audit trail of employee career progression. Created whenever Employee.position_id, department_id, employment_type_id, or salary changes. Enables: (1) accurate historical payroll (use salary snapshot), (2) tenure calculation by status periods, (3) promotion/transfer reporting. The effective_date/end_date range defines the period this record applies to; NULL end_date means current. remarks capture business context for changes.',
                   jh_fields,
                   relationships=['Many-to-One with Employee', 'Optional FKs to Position/EmploymentType/Department'],
                   notes='• Never update historical records; insert new record for changes\n• Validate that effective_date > previous record\'s effective_date\n• When creating new record, set previous record\'s end_date = new effective_date - 1 day\n• Consider adding approved_by field for change authorization')
    
    # 26. TAX
    tax_fields = [
        {'name': 'id', 'type': 'Integer', 'null': 'No', 'default': 'Auto', 'description': 'Primary key for tax bracket definition'},
        {'name': 'min_income', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Minimum taxable income for this bracket (inclusive)'},
        {'name': 'max_income', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Maximum taxable income for this bracket (inclusive); NULL for top bracket'},
        {'name': 'rate', 'type': 'Float', 'null': 'Yes', 'default': 'NULL', 'description': 'Marginal tax rate percentage (e.g., 0.20 for 20%)'},
        {'name': 'fixed', 'type': 'Float', 'null': 'Yes', 'default': '0', 'description': 'Base tax amount for progressive calculation: tax = fixed + (income × rate)'},
    ]
    add_model_table(doc, '26. Tax', 'tax',
                   'Philippine income tax brackets for withholding tax calculation. Implements TRAIN Law progressive tax schedule: for income in [min_income, max_income], tax = fixed + (income × rate). The compute(income) method finds the matching bracket and applies the formula. Used during Payroll.calculate() to determine withholding tax deduction. Update brackets when BIR issues new tax tables; historical payroll uses rules at time of calculation.',
                   tax_fields,
                   relationships=['Used by Deduction with calculation_type="progressive" for tax deductions'],
                   notes='• Brackets should cover all income ranges without gaps or overlaps\n• max_income=NULL indicates the top bracket (no upper limit)\n• fixed amount represents tax on lower brackets in progressive system\n• Consider adding effective_date for time-based tax law changes')
    
    # ===== FINALIZE DOCUMENT =====
    doc.add_page_break()
    footer = doc.add_paragraph('End of Data Dictionary')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True
    
    # Save the document
    output_path = 'GOVHRPAY_Data_Dictionary.docx'
    doc.save(output_path)
    print(f"✓ Data dictionary created successfully: {output_path}")
    print(f"✓ Contains 26 model tables, each on its own page")
    print(f"✓ Tables formatted to prevent splitting across pages")
    return output_path

if __name__ == "__main__":
    create_data_dictionary()