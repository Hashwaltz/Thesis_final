import io
import os
from io import BytesIO
import pandas as pd
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from flask import current_app, send_file
from datetime import datetime
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image,
    Table,
    TableStyle
)
from reportlab.lib.pagesizes import A4


from main_app.models.hr_models import Employee
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib import colors



def generate_moa_excel(
    employees_by_type,
    agency_name="LGU-NORZAGARAY, BULACAN",
    regional_office="3",
    report_prefix="LIST OF"
):
    
    wb = Workbook()

    bold_center = Font(bold=True)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )
    gray_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

    for i, (etype, employees) in enumerate(employees_by_type.items()):
        ws = wb.active if i == 0 else wb.create_sheet()
        ws.title = etype.name[:28]

        # ===== MAIN HEADER =====
        ws.merge_cells("A1:K1")
        ws["A1"] = f"{report_prefix} {etype.name.upper()} PERSONNEL"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A1"].alignment = Alignment(horizontal="center")

        ws.merge_cells("A2:K2")
        ws["A2"] = f"(As of {date.today().strftime('%B %d, %Y')})"
        ws["A2"].font = Font(italic=True)
        ws["A2"].alignment = Alignment(horizontal="center")

        # ===== AGENCY INFO =====
        ws.append([])
        ws.append(["Agency name:", agency_name])
        ws.append(["Regional Office No:", regional_office])
        ws.append([])

        start_row = ws.max_row + 1

        # ===== TWO-LEVEL HEADER =====
        ws.merge_cells(f"A{start_row}:A{start_row+1}")
        ws.merge_cells(f"B{start_row}:D{start_row}")
        ws.merge_cells(f"E{start_row}:E{start_row+1}")
        ws.merge_cells(f"F{start_row}:F{start_row+1}")
        ws.merge_cells(f"G{start_row}:G{start_row+1}")
        ws.merge_cells(f"H{start_row}:H{start_row+1}")
        ws.merge_cells(f"I{start_row}:I{start_row+1}")
        ws.merge_cells(f"J{start_row}:K{start_row}")

        ws[f"A{start_row}"] = "NO."
        ws[f"B{start_row}"] = "Name of Personnel"
        ws[f"E{start_row}"] = "DATE OF BIRTH\n(MM/DD/YYYY)"
        ws[f"F{start_row}"] = "SEX\n(pls. select)"
        ws[f"G{start_row}"] = "Level of CS Eligibility\n(pls. select)"
        ws[f"H{start_row}"] = "WORK STATUS\n(pls. select)"
        ws[f"I{start_row}"] = f"No. of Years of Service as {etype.name} personnel"
        ws[f"J{start_row}"] = "NATURE OF WORK"

        ws[f"B{start_row+1}"] = "SURNAME"
        ws[f"C{start_row+1}"] = "FIRST NAME/\nEXTENSION NAME"
        ws[f"D{start_row+1}"] = "MIDDLE INITIAL"
        ws[f"J{start_row+1}"] = "Pls select"
        ws[f"K{start_row+1}"] = "Please specify"

        # ===== STYLE HEADER =====
        for row in ws.iter_rows(min_row=start_row, max_row=start_row+1, min_col=1, max_col=11):
            for cell in row:
                cell.font = bold_center
                cell.alignment = center_align
                cell.border = thin_border
                cell.fill = gray_fill

        # ===== TABLE BODY =====
        for idx, emp in enumerate(employees, start=1):
            ws.append([
                idx,
                emp.last_name or "",
                emp.first_name or "",
                emp.middle_name or "",
                emp.date_of_birth.strftime("%m/%d/%Y") if emp.date_of_birth else "",
                emp.gender or "",
                getattr(emp, "cs_eligibility", "No eligibility"),
                emp.status or "Active",
                emp.get_working_duration() if hasattr(emp, "get_working_duration") else "",
                emp.position.name if emp.position else "",
                "",
            ])

        # ===== AUTO WIDTH =====
        for i_col, col in enumerate(ws.columns, start=1):
            max_length = max((len(str(cell.value)) for cell in col if cell.value), default=0)
            ws.column_dimensions[get_column_letter(i_col)].width = max_length + 2

    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    return file_stream


def generate_excel_employees(
    data,
    headers,
    title="REPORT",
    agency_name="LGU-NORZAGARAY, BULACAN",
    regional_office="3"
):
    """
    Universal Excel Report Generator

    data → list of lists
    headers → list of column headers
    """

    wb = Workbook()
    ws = wb.active

    bold_center = Font(bold=True)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    gray_fill = PatternFill(
        start_color="D9D9D9",
        end_color="D9D9D9",
        fill_type="solid"
    )

    # ===== HEADER =====
    ws.merge_cells(f"A1:{get_column_letter(len(headers))}1")
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells(f"A2:{get_column_letter(len(headers))}2")
    ws["A2"] = f"(As of {date.today().strftime('%B %d, %Y')})"
    ws["A2"].alignment = Alignment(horizontal="center")

    ws.append([])
    ws.append(["Agency name:", agency_name])
    ws.append(["Regional Office No:", regional_office])
    ws.append([])

    start_row = ws.max_row + 1

    # ===== TABLE HEADER =====
    ws.append(headers)

    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=start_row, column=col)
        cell.font = bold_center
        cell.alignment = center_align
        cell.border = thin_border
        cell.fill = gray_fill

    # ===== BODY DATA =====
    for row in data:
        ws.append(row)

    # ===== AUTO WIDTH =====
    for i_col, col in enumerate(ws.columns, start=1):
        max_length = max(
            (len(str(cell.value)) for cell in col if cell.value),
            default=0
        )
        ws.column_dimensions[get_column_letter(i_col)].width = max_length + 2

    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    return file_stream

def generate_service_record_docx(employee):
    doc = Document()

    # ================= HEADER =================
    p1 = doc.add_paragraph("Republic of the Philippines")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph("NORZAGARAY, REGION 3")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True

    title = doc.add_paragraph("S E R V I C E   R E C O R D")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph("")

    # ================= NAME SECTION =================
    doc.add_paragraph("Name :")
    table_name = doc.add_table(rows=1, cols=3)
    table_name.style = "Table Grid"
    table_name.rows[0].cells[0].text = "First Name"
    table_name.rows[0].cells[1].text = "Middle Name"
    table_name.rows[0].cells[2].text = "Last Name"

    row = table_name.add_row().cells
    row[0].text = employee.first_name or ""
    row[1].text = employee.middle_name or ""
    row[2].text = employee.last_name or ""

    doc.add_paragraph("(If married woman, give full maiden name)")

    birth = employee.date_of_birth.strftime("%B %d, %Y") if employee.date_of_birth else ""
    doc.add_paragraph(f"Date and place of birth : {birth}")
    doc.add_paragraph("(Date herein should be checked from birth or baptismal certificate or some other reliable documents)")
    doc.add_paragraph("B.P. Number: __________________")
    doc.add_paragraph("TIN # : __________________")
    doc.add_paragraph("")

    # ================= CERTIFICATION =================
    cert = doc.add_paragraph(
        "This is to certify that the employee named hereunder actually rendered services in this Office as shown by the "
        "service record below, each line of which is supported by appointment and other papers actually issued by this "
        "Office and approved by the authorities concerned."
    )
    cert.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    # ================= SERVICE RECORD TABLE =================
    table = doc.add_table(rows=2, cols=10)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ---- TOP HEADER ----
    top = table.rows[0].cells
    top[0].text = "SERVICE (Inclusive Dates)"
    top[1].text = ""  # merged later
    top[2].text = "RECORD OF APPOINTMENT"
    top[3].text = "OFFICE / ENTITY / DIVISION"
    top[4].text = "Leave(s)"
    top[5].text = ""
    top[6].text = "SEPARATION"
    top[7].text = ""
    top[8].text = "Remarks"
    top[9].text = ""

    # Merge top row cells as per format
    top[0].merge(top[1])
    top[4].merge(top[5])
    top[6].merge(top[7])

    # ---- SECOND HEADER ----
    second = table.rows[1].cells
    second[0].text = "From"
    second[1].text = "To"
    second[2].text = "Designation / Status (1)"
    second[3].text = "Station / Place of Assignment"
    second[4].text = "With Pay"
    second[5].text = "Without Pay"
    second[6].text = "Date"
    second[7].text = "Cause"
    second[8].text = ""
    second[9].text = ""

    # ================= JOB HISTORY DATA =================
    histories = sorted(employee.job_history, key=lambda x: x.effective_date)

    for h in histories:
        row = table.add_row().cells
        row[0].text = h.effective_date.strftime("%b %d, %Y") if h.effective_date else ""
        row[1].text = h.end_date.strftime("%b %d, %Y") if h.end_date else "Present"

        designation = h.position.name if h.position else ""
        if h.employment_type:
            designation += f" ({h.employment_type.name})"
        row[2].text = designation

        row[3].text = h.department.name if h.department else ""
        row[4].text = ""
        row[5].text = ""
        row[6].text = ""
        row[7].text = ""
        row[8].text = h.remarks or ""
        row[9].text = ""

    doc.add_paragraph("")

    # ================= ISSUED STATEMENT =================
    issue = doc.add_paragraph(
        "Issued on compliance with Executive Order No. 54 dated August 10, 1954, and in accordance with Circular No. 68 dated August 10, 1954 of the System."
    )
    issue.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    # ================= SIGNATURE =================
    doc.add_paragraph("CERTIFIED CORRECT:")
    doc.add_paragraph("")
    doc.add_paragraph("FERNANDO DG. CRUZ")
    doc.add_paragraph("Acting MHRMO")
    doc.add_paragraph("")

    # ================= DATE + PAGE =================
    doc.add_paragraph(date.today().strftime("%A, %B %d, %Y"))
    doc.add_paragraph("Page 1 of 1")

    # ================= FONT STANDARD =================
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer




def generate_coe_pdf(employee, fields=None):
    """
    Generate COE PDF with conditional field inclusion.
    `fields`: list of fields to include ['position', 'department', 'employment_type', 
             'hire_date', 'end_date', 'working_duration', 'deductions']
    """
    fields = fields or []
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=inch, leftMargin=inch,
        topMargin=inch, bottomMargin=inch
    )
    
    styles = getSampleStyleSheet()
    OFFICIAL_BLUE = HexColor('#000080')

    # ─── STYLES ────────────────────────────────────────────────────────────────
    header_style = ParagraphStyle(
        'Header', parent=styles['Normal'],
        alignment=TA_CENTER, fontSize=10, leading=12, textColor=OFFICIAL_BLUE
    )
    header_bold_style = ParagraphStyle(
        'HeaderBold', parent=header_style,
        fontSize=11, fontName='Helvetica-Bold'
    )
    date_style = ParagraphStyle(
        'Date', parent=styles['Normal'],
        alignment=TA_RIGHT, fontSize=11, spaceAfter=20
    )
    title_style = ParagraphStyle(
        'Title', parent=styles['Normal'],
        alignment=TA_CENTER, fontSize=14, fontName='Helvetica-Bold',
        spaceAfter=20
    )
    body_style = ParagraphStyle(
        'Body', parent=styles['Normal'],
        alignment=TA_LEFT, fontSize=11, leading=18,
        firstLineIndent=0.5*inch, spaceAfter=12
    )
    signature_style = ParagraphStyle(
        'Signature', parent=styles['Normal'],
        alignment=TA_RIGHT, fontSize=11, leading=14, spaceBefore=30
    )

    story = []

    # ─── LOGO ──────────────────────────────────────────────────────────────────
    logo_path = os.path.join(current_app.root_path, "static", "img", "garay.png")
    if os.path.exists(logo_path):
        logo = Image(logo_path, width=0.85*inch, height=0.85*inch)
        story.append(Table([[logo]], colWidths=[6.5*inch], style=[('ALIGN', (0,0), (-1,-1), 'CENTER')]))
        story.append(Spacer(1, 6))

    # ─── HEADER ───────────────────────────────────────────────────────────────
    story.append(Paragraph("Republic of the Philippines", header_style))
    story.append(Paragraph("Province of Bulacan", header_style))
    story.append(Paragraph("MUNICIPALITY OF NORZAGARAY", header_bold_style))
    story.append(Paragraph("HUMAN RESOURCE MANAGEMENT OFFICE", header_bold_style))

    # ─── BLUE SEPARATOR LINE ───────────────────────────────────────────────────
    separator = Table([[""]], colWidths=[6.5*inch])
    separator.setStyle(TableStyle([
        ('LINEBELOW', (0,0), (-1,-1), 4, OFFICIAL_BLUE),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4)
    ]))
    story.append(separator)
    story.append(Spacer(1, 16))

    # ─── DATE ──────────────────────────────────────────────────────────────────
    today_date = datetime.now().strftime("%B %d, %Y")
    story.append(Paragraph(today_date, date_style))

    # ─── TITLE ─────────────────────────────────────────────────────────────────
    story.append(Paragraph("CERTIFICATION", title_style))

    # ─── SALUTATION ────────────────────────────────────────────────────────────
    story.append(Paragraph("<b>TO WHOM IT MAY CONCERN:</b>", body_style))
    story.append(Spacer(1, 8))

    # ─── DYNAMIC DATA EXTRACTION ───────────────────────────────────────────────
    gender = (getattr(employee, 'gender', '') or '').lower()
    prefix = "Mr." if gender == "male" else ("Ms." if gender == "female" else "Mr./Ms.")
    
    first_name = getattr(employee, 'first_name', '') or ''
    middle_name = getattr(employee, 'middle_name', '') or ''
    last_name = getattr(employee, 'last_name', '') or ''
    full_name = f"{first_name} {middle_name} {last_name}".strip().upper()

    # Conditional data fetching based on fields
    position = getattr(getattr(employee, 'position', None), 'name', None) if 'position' in fields else None
    department = getattr(getattr(employee, 'department', None), 'name', None) if 'department' in fields else None
    emp_type = getattr(getattr(employee, 'employment_type', None), 'name', None) if 'employment_type' in fields else None
    
    hire_date = None
    if 'hire_date' in fields and getattr(employee, 'date_hired', None):
        hire_date = employee.date_hired.strftime("%B %d, %Y")
    
    end_date = None
    if 'end_date' in fields:
        if getattr(employee, 'status', '') == 'Active':
            end_date = "Present"
        elif getattr(employee, 'end_date', None):
            end_date = employee.end_date.strftime("%B %d, %Y")
    
    working_duration = None
    if 'working_duration' in fields and hasattr(employee, 'get_working_duration'):
        working_duration = employee.get_working_duration()

    # ─── BUILD BODY TEXT WITH CONDITIONAL FIELDS ───────────────────────────────
    body_parts = [f"This is to certify that <b>{full_name}</b> had been employed in this Local Government Unit, detailed at the Norzagaray College"]
    
    # Add position if requested
    if position:
        body_parts.append(f"as <b>{position}</b>")
    
    # Add department if requested (with proper connector)
    if department:
        if position:
            body_parts.append(f"under the <b>{department}</b>")
        else:
            body_parts.append(f"under the <b>{department}</b>")
    
    # Add employment type if requested
    if emp_type:
        body_parts.append(f"on a <b>{emp_type}</b> Status")
    
    # Add date range if any date fields are requested
    if hire_date or end_date:
        date_clause = "since" if hire_date else ""
        if hire_date:
            date_clause += f" <b>{hire_date}</b>"
        if end_date:
            date_clause += f" up to <b>{end_date}</b>"
        body_parts.append(date_clause.strip())
    
    # Add working duration if requested
    if working_duration:
        body_parts.append(f"with a total duration of <b>{working_duration}</b>")
    
    # Join body parts with proper punctuation
    body_text = " ".join(body_parts).rstrip()
    if not body_text.endswith('.'):
        body_text += "."
    
    story.append(Paragraph(body_text, body_style))

    # ─── ADD DEDUCTIONS IF REQUESTED ───────────────────────────────────────────
    if 'deductions' in fields:
        deductions = []
        for ed in getattr(employee, "employee_deductions", []):
            deduction_name = getattr(getattr(ed, 'deduction', None), 'name', None)
            if deduction_name:
                deductions.append(deduction_name)
        if deductions:
            ded_text = f"Deductions: {', '.join(deductions)}."
            story.append(Paragraph(ded_text, body_style))

    # ─── ISSUED STATEMENT ──────────────────────────────────────────────────────
    issued_text = f"Issued upon request of {prefix} {last_name} for whatever legal purpose this may serve."
    story.append(Paragraph(issued_text, body_style))

    # ─── SIGNATURE BLOCK ───────────────────────────────────────────────────────
    sig_text = "<b>FERNANDO DG. CRUZ</b><br/>Acting MHRMO"
    story.append(Paragraph(sig_text, signature_style))

    # ─── BUILD PDF ─────────────────────────────────────────────────────────────
    doc.build(story)
    buffer.seek(0)
    return buffer




def safe_get(obj, attr, default=0):
    return getattr(obj, attr, default) or default


def export_payroll_excel(query):

    payrolls = query.all()
    data = []

    for p in payrolls:

        emp = p.employee

        # ===============================
        # DEDUCTION BREAKDOWN (REAL DATA)
        # ===============================
        sss = philhealth = pagibig = other = 0

        for d in getattr(p, "deduction_breakdown", []):
            name = (d.deduction_name or "").lower()

            if "sss" in name:
                sss += d.employee_share or 0
            elif "philhealth" in name:
                philhealth += d.employee_share or 0
            elif "pag" in name:
                pagibig += d.employee_share or 0
            else:
                other += d.employee_share or 0

        # ===============================
        # ALLOWANCES
        # ===============================
        total_allowances = p.allowance_total or 0

        # ===============================
        # ATTENDANCE (SAFE)
        # ===============================
        try:
            days_worked = p.compute_attendance_days()
        except:
            days_worked = 0

        try:
            hours_worked = p.compute_attendance_hours()
        except:
            hours_worked = 0

        # ===============================
        # DATA ROW
        # ===============================
        data.append({

            "Employee ID": emp.employee_id,
            "Employee Name": emp.get_full_name(),
            "Department": emp.department.name if emp.department else "-",
            "Employment Type": emp.employment_type.name if emp.employment_type else "-",

            # Attendance
            "Days Worked": days_worked,
            "Hours Worked": hours_worked,

            # Earnings
            "Basic Salary": p.basic_salary or 0,
            "Overtime Pay": p.overtime_pay or 0,
            "Holiday Pay": p.holiday_pay or 0,
            "Night Differential": p.night_diff or 0,
            "Allowances": total_allowances,

            "Gross Pay": p.gross_pay or 0,

            # Deductions
            "SSS": sss,
            "PhilHealth": philhealth,
            "Pag-IBIG": pagibig,
            "Tax Withheld": getattr(p, "tax_withheld", 0),
            "Other Deductions": other,

            "Total Deductions": p.total_deductions or 0,

            # Net
            "Net Pay": p.net_pay or 0,

            # Payroll Info
            "Status": p.status or "-",
            "Payslip No": p.payslip.payslip_number if p.payslip else "-",
            "Pay Date": p.period.pay_date if p.period else "-",
            "Pay Period": f"{p.period.start_date} - {p.period.end_date}" if p.period else "-",

            # HR Insight
            "Years of Service": emp.years_of_service
        })

    # ===============================
    # DATAFRAME
    # ===============================
    df = pd.DataFrame(data)

    # Sort for better readability
    df = df.sort_values(by=["Department", "Employee Name"])

    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:

        df.to_excel(
            writer,
            index=False,
            sheet_name="Payroll",
            startrow=12
        )

        workbook = writer.book
        worksheet = writer.sheets["Payroll"]

        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.drawing.image import Image as OpenpyxlImage

        # ==========================================
        # LOGO
        # ==========================================
        logo_path = r"C:\Users\pc\Desktop\Thesis_final\main_app\static\img\garay.png"

        if os.path.exists(logo_path):
            logo = OpenpyxlImage(logo_path)
            logo.width = 110
            logo.height = 110
            worksheet.add_image(logo, "A1")

        # ==========================================
        # HEADER
        # ==========================================
        worksheet.merge_cells("A2:R2")
        worksheet["A2"] = "Republic of the Philippines"

        worksheet.merge_cells("A3:R3")
        worksheet["A3"] = "MUNICIPALITY OF NORZAGARAY"

        worksheet.merge_cells("A4:R4")
        worksheet["A4"] = "Province of Bulacan"

        worksheet.merge_cells("A6:R6")
        worksheet["A6"] = "Municipal Hall of Norzagaray"

        worksheet.merge_cells("A7:R7")
        worksheet["A7"] = "Norzagaray, Bulacan"

        worksheet.merge_cells("A9:R9")
        worksheet["A9"] = "PAYROLL SUMMARY REPORT"

        header_font = Font(size=12, bold=True)
        title_font = Font(size=16, bold=True)

        for cell in ["A2","A3","A4","A6","A7"]:
            worksheet[cell].alignment = Alignment(horizontal="center")
            worksheet[cell].font = header_font

        worksheet["A9"].alignment = Alignment(horizontal="center")
        worksheet["A9"].font = title_font

        # ==========================================
        # SUMMARY (NEW 🔥)
        # ==========================================
        worksheet["A10"] = f"Total Employees: {len(payrolls)}"
        worksheet["A11"] = f"Total Payroll: ₱{df['Net Pay'].sum():,.2f}"

        # ==========================================
        # HEADER STYLE
        # ==========================================
        header_row = 13

        fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

        for cell in worksheet[header_row]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")
            cell.fill = fill

        # ==========================================
        # FREEZE + FILTER 🔥
        # ==========================================
        worksheet.freeze_panes = "A14"
        worksheet.auto_filter.ref = f"A13:R{worksheet.max_row}"

        # ==========================================
        # BORDERS
        # ==========================================
        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for row in worksheet.iter_rows(
            min_row=header_row,
            max_row=worksheet.max_row,
            min_col=1,
            max_col=worksheet.max_column
        ):
            for cell in row:
                cell.border = border

        # ==========================================
        # CURRENCY FORMAT
        # ==========================================
        currency_columns = [
            "G","H","I","J","K","L","M","N","O","P"
        ]

        for col in currency_columns:
            for row in range(header_row+1, worksheet.max_row+1):
                worksheet[f"{col}{row}"].number_format = '₱#,##0.00'

        # ==========================================
        # AUTO WIDTH
        # ==========================================
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter

            for cell in column:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass

            worksheet.column_dimensions[column_letter].width = max_length + 3

    output.seek(0)

    filename = f"Payroll_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ======================================================
# PAYROLL SUMMARY REPORT
# ======================================================

def payroll_summary_report(payrolls):

    data = []

    for p in payrolls:

        data.append({
            "Employee ID": p.employee.employee_id,
            "Employee Name": f"{p.employee.first_name} {p.employee.last_name}",
            "Department": p.employee.department.name if p.employee.department else "-",

            "Basic Salary": p.basic_salary,
            "Overtime Pay": p.overtime_pay,
            "Holiday Pay": p.holiday_pay,
            "Night Differential": p.night_diff,
            "Gross Pay": p.gross_pay,

            "Total Deductions": p.total_deductions,
            "Net Pay": p.net_pay,

            "Payroll Period":
            f"{p.period.start_date} - {p.period.end_date}"
        })

    df = pd.DataFrame(data)

    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Payroll Summary")

    output.seek(0)

    filename = f"Payroll_Summary_{datetime.now().strftime('%Y%m%d')}.xlsx"

    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ======================================================
# DEDUCTION SUMMARY REPORT
# ======================================================

def deduction_summary_report(payrolls):

    deductions = {}

    for p in payrolls:

        for d in p.deduction_breakdown:

            name = d.deduction_name

            deductions.setdefault(name, 0)

            deductions[name] += d.employee_share

    data = [
        {"Deduction": k, "Total": v}
        for k, v in deductions.items()
    ]

    df = pd.DataFrame(data)

    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Deduction Summary")

    output.seek(0)

    filename = f"Deduction_Summary_{datetime.now().strftime('%Y%m%d')}.xlsx"

    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ======================================================
# EMPLOYEE PAYROLL HISTORY
# ======================================================

def employee_payroll_history(employee):

    data = []

    for p in employee.payrolls:

        data.append({
            "Payroll Period": f"{p.period.start_date} - {p.period.end_date}",
            "Gross Pay": p.gross_pay,
            "Total Deductions": p.total_deductions,
            "Net Pay": p.net_pay,
            "Status": p.status
        })

    df = pd.DataFrame(data)

    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Payroll History")

    output.seek(0)

    filename = f"Payroll_History_{employee.employee_id}.xlsx"

    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )






# ======================================================
# EMPLOYEES BY YEARS OF SERVICE REPORT
# ======================================================

def export_employees_by_year_of_service(employees):
    wb = Workbook()
    ws = wb.active
    ws.title = "Employees"

    # --- HEADER SECTION ---
    ws.merge_cells("A1:H1")
    ws["A1"] = "LIST OF PERSONNEL"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:H2")
    ws["A2"] = f"(As of {date.today().strftime('%B %d, %Y')})"
    ws["A2"].alignment = Alignment(horizontal="center")

    ws["A4"] = "Agency name:"
    ws["B4"] = "LGU-NORZAGARAY, BULACAN"

    ws["A5"] = "Regional Office No:"
    ws["B5"] = "3"

    # --- TABLE HEADER ---
    headers = [
        "Full Name",
        "Email",
        "Phone",
        "Department",
        "Position",
        "Date Hired",
        "Years of Service",
        "Barangay"
    ]

    ws.append([])  # blank row
    ws.append(headers)

    # Apply bold font to header row
    header_row = ws.max_row
    for cell in ws[header_row]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # --- EMPLOYEE DATA ---
    for e in employees:
        ws.append([
            e.get_full_name(),
            e.email,
            e.phone or "-",
            e.department.name if e.department else "-",
            e.position.name if e.position else "-",
            e.date_hired.strftime("%Y-%m-%d") if e.date_hired else "-",
            e.get_working_duration(),
            e.barangay or "-"
        ])

    # --- Auto-fit column widths safely ---
    for column_cells in ws.columns:
        max_length = 0
        col_letter = get_column_letter(column_cells[0].column)
        for cell in column_cells:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = max_length + 2

    # --- Save to in-memory buffer ---
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="Employee_Report.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )




def generate_payslip_pdf(payroll):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.8*cm,
        leftMargin=1.8*cm,
        topMargin=2.0*cm,
        bottomMargin=1.5*cm
    )
    
    styles = getSampleStyleSheet()
    
    # 🎨 Custom Styles
    title_style = ParagraphStyle('Title', parent=styles['Normal'],
        fontSize=18, textColor=colors.HexColor('#111827'), fontName='Helvetica-Bold',
        alignment=TA_CENTER, spaceAfter=4)
    
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'],
        fontSize=10, textColor=colors.HexColor('#6b7280'), alignment=TA_CENTER, spaceAfter=12)
    
    section_style = ParagraphStyle('Section', parent=styles['Normal'],
        fontSize=11, textColor=colors.HexColor('#0f172a'), fontName='Helvetica-Bold',
        spaceBefore=16, spaceAfter=8)
    
    label_style = ParagraphStyle('Label', parent=styles['Normal'],
        fontSize=8, textColor=colors.HexColor('#6b7280'), spaceAfter=2)
    
    value_style = ParagraphStyle('Value', parent=styles['Normal'],
        fontSize=10, textColor=colors.HexColor('#111827'), spaceAfter=6)
    
    money_style = ParagraphStyle('Money', parent=value_style,
        fontName='Helvetica', alignment=TA_RIGHT, spaceAfter=6)
    
    total_green = ParagraphStyle('TotalGreen', parent=money_style,
        fontName='Helvetica-Bold', fontSize=11, textColor=colors.HexColor('#15803d'))
    
    total_red = ParagraphStyle('TotalRed', parent=money_style,
        fontName='Helvetica-Bold', fontSize=11, textColor=colors.HexColor('#b91c1c'))
    
    net_pay_style = ParagraphStyle('NetPay', parent=styles['Normal'],
        fontSize=20, textColor=colors.HexColor('#15803d'), fontName='Helvetica-Bold',
        alignment=TA_RIGHT, spaceAfter=0)
    
    sig_style = ParagraphStyle('Signature', parent=styles['Normal'],
        fontSize=9, textColor=colors.HexColor('#374151'), alignment=TA_CENTER, spaceBefore=2)
    
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'],
        fontSize=7, textColor=colors.HexColor('#9ca3af'), alignment=TA_CENTER, spaceBefore=10)

    elements = []

    # ── TOP ACCENT BAR ──
    header_table = Table([['']], colWidths=[A4[0] - 3.6*cm])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#15803d')),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 8))

    # ── HEADER ──
    elements.append(Paragraph("MUNICIPALITY OF NORZAGARAY", title_style))
    elements.append(Paragraph("Municipal Human Resource Management Office", subtitle_style))
    
    # Badge
    badge_table = Table([["OFFICIAL PAYSLIP"]], colWidths=[3.5*cm], hAlign='RIGHT')
    badge_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#dcfce7')),
        ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor('#15803d')),
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('ALIGNMENT', (0,0), (-1,-1), 'CENTER'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#86efac')),
    ]))
    elements.append(badge_table)
    elements.append(Spacer(1, 12))

    # ── EMPLOYEE INFO ──
    emp = payroll.employee
    period = payroll.period
    
    info_data = [
        [Paragraph("EMPLOYEE INFORMATION", section_style), '', ''],
        [Paragraph("Full Name", label_style), Paragraph("Position", label_style), Paragraph("Pay Date", label_style)],
        [Paragraph(emp.get_full_name(), value_style), Paragraph(emp.position.name, value_style), 
         Paragraph(period.pay_date.strftime('%b %d, %Y') if period.pay_date else '-', value_style)],
        [Paragraph("Department", label_style), Paragraph("Employment Type", label_style), Paragraph("Coverage Period", label_style)],
        [Paragraph(emp.department.name, value_style), Paragraph(emp.employment_type.name, value_style), 
         Paragraph(f"{period.start_date.strftime('%m/%d/%y')} - {period.end_date.strftime('%m/%d/%y')}", value_style)],
    ]
    
    info_table = Table(info_data, colWidths=[5.5*cm, 5.5*cm, 3.5*cm])
    info_table.setStyle(TableStyle([
        ('LINEBELOW', (0,0), (-1,0), 1, colors.HexColor('#e5e7eb')),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f8fafc')),
        ('TOPPADDING', (0,1), (-1,-1), 4),
        ('BOTTOMPADDING', (0,1), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 16))

    # ── EARNINGS ──
    elements.append(Paragraph("EARNINGS", section_style))
    
    earnings_data = [["Description", "Amount"]]
    earnings_data.append(["Basic Pay", f"₱{(payroll.gross_pay - payroll.allowance_total or 0):,.2f}"])
    earnings_data.append(["Allowance", f"₱{payroll.allowance_total or 0:,.2f}"])
    earnings_data.append(["Overtime Pay", f"₱{payroll.overtime_pay or 0:,.2f}"])
    earnings_data.append([Paragraph("<b>TOTAL EARNINGS</b>", section_style), Paragraph(f"₱{payroll.gross_pay or 0:,.2f}", total_green)])
    
    earnings_table = Table(earnings_data, colWidths=[11*cm, 3.5*cm])
    earnings_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0,0), (0,0), colors.HexColor('#6b7280')),
        ('TEXTCOLOR', (1,0), (1,0), colors.HexColor('#6b7280')),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 8),
        ('LINEBELOW', (0,0), (-1,0), 1, colors.HexColor('#d1d5db')),
        ('LINEBELOW', (0,1), (-1,-2), 0.5, colors.HexColor('#e5e7eb')),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#f0fdf4')),
        ('LINEABOVE', (0,-1), (-1,-1), 1, colors.HexColor('#86efac')),
        ('ALIGN', (1,1), (1,-1), 'RIGHT'),
        ('TOPPADDING', (0,1), (-1,-1), 5),
        ('BOTTOMPADDING', (0,1), (-1,-1), 5),
        ('LEFTPADDING', (0,1), (-1,-1), 8),
    ]))
    elements.append(earnings_table)
    elements.append(Spacer(1, 14))

    # ── DEDUCTIONS ──
    elements.append(Paragraph("DEDUCTIONS", section_style))
    
    deductions_data = [["Description", "Amount"]]
    for d in payroll.deduction_breakdown:
        deductions_data.append([d.deduction_name, f"₱{d.employee_share or 0:,.2f}"])
    for l in payroll.loan_payments:
        deductions_data.append([f"{l.loan.provider} - {l.loan.loan_type}", f"₱{l.amount_paid or 0:,.2f}"])
    
    if not payroll.deduction_breakdown and not payroll.loan_payments:
        deductions_data.append(["No deductions for this period", ""])
        
    deductions_data.append([Paragraph("<b>TOTAL DEDUCTIONS</b>", section_style), Paragraph(f"₱{payroll.total_deductions or 0:,.2f}", total_red)])
    
    deductions_table = Table(deductions_data, colWidths=[11*cm, 3.5*cm])
    deductions_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0,0), (0,0), colors.HexColor('#6b7280')),
        ('TEXTCOLOR', (1,0), (1,0), colors.HexColor('#6b7280')),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 8),
        ('LINEBELOW', (0,0), (-1,0), 1, colors.HexColor('#d1d5db')),
        ('LINEBELOW', (0,1), (-1,-2), 0.5, colors.HexColor('#e5e7eb')),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#fef2f2')),
        ('LINEABOVE', (0,-1), (-1,-1), 1, colors.HexColor('#fca5a5')),
        ('ALIGN', (1,1), (1,-1), 'RIGHT'),
        ('TOPPADDING', (0,1), (-1,-1), 5),
        ('BOTTOMPADDING', (0,1), (-1,-1), 5),
        ('LEFTPADDING', (0,1), (-1,-1), 8),
    ]))
    elements.append(deductions_table)
    elements.append(Spacer(1, 20))

    # ── NET PAY ──
    net_table = Table([["NET TAKE-HOME PAY", f"₱{payroll.net_pay or 0:,.2f}"]], colWidths=[9*cm, 5.5*cm])
    net_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f0fdf4')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#86efac')),
        ('TOPPADDING', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING', (0,0), (-1,-1), 12),
        ('FONTNAME', (0,0), (0,0), 'Helvetica'),
        ('FONTSIZE', (0,0), (0,0), 10),
        ('TEXTCOLOR', (0,0), (0,0), colors.HexColor('#374151')),
        ('ALIGNMENT', (0,0), (0,0), 'LEFT'),
        ('ALIGNMENT', (1,0), (1,0), 'RIGHT'),
        ('FONTNAME', (1,0), (1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (1,0), (1,0), 18),
        ('TEXTCOLOR', (1,0), (1,0), colors.HexColor('#15803d')),
    ]))
    elements.append(net_table)
    elements.append(Spacer(1, 30))

    # ── SIGNATURES ──
    sig_data = [
        [
            Paragraph("Prepared By:", sig_style),
            Paragraph("Noted By:", sig_style)
        ],
        [
            Paragraph("__________________________", sig_style),
            Paragraph("__________________________", sig_style)
        ],
        [
            Paragraph("HR Payroll Staff", sig_style),
            Paragraph("HR Manager", sig_style)
        ],
        [
            Paragraph("Municipal HRMO", sig_style),
            Paragraph("Municipal HRMO", sig_style)
        ]
    ]
    
    sig_table = Table(sig_data, colWidths=[7*cm, 7*cm])
    sig_table.setStyle(TableStyle([
        ('TOPPADDING', (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
    ]))
    elements.append(sig_table)
    
    # ── FOOTER ──
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')} • Municipality of Norzagaray HRMO • Confidential",
        footer_style
    ))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer