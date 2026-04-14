from flask import  render_template, request, flash, redirect, url_for,  send_file
from io import BytesIO 
from collections import Counter, defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_login import login_required, current_user
from datetime import datetime, date, timedelta
import logging
from sqlalchemy import func

from main_app.helpers.decorators import admin_required  
from main_app.models.hr_models import  Employee, Department, Attendance, Leave, LeaveType


from main_app.blueprints.hr_system.routes.admin import hr_admin_bp

logger = logging.getLogger(__name__)


# ------------------------- Reports -------------------------
@hr_admin_bp.route('/reports')
@login_required
@admin_required
def reports():
    # Get filters from query parameters
    report_type = request.args.get('report_type', 'attendance')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    page = request.args.get('page', 1, type=int)

    # Convert dates to datetime objects if provided
    try:
        start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date() if start_date else None
        end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date() if end_date else None
    except ValueError:
        flash("Invalid date format", "error")
        start_date_obj = end_date_obj = None

    # Fetch data based on report type
    if report_type == 'attendance':
        query = Attendance.query
        if start_date_obj:
            query = query.filter(Attendance.date >= start_date_obj)
        if end_date_obj:
            query = query.filter(Attendance.date <= end_date_obj)
        data = query.order_by(Attendance.date.desc()).paginate(page=page, per_page=20)

        employees = Employee.query.filter_by(active=True).all()  # for filter dropdown

    elif report_type == 'leaves':
        query = Leave.query
        if start_date_obj:
            query = query.filter(Leave.start_date >= start_date_obj)
        if end_date_obj:
            query = query.filter(Leave.end_date <= end_date_obj)
        data = query.order_by(Leave.start_date.desc()).paginate(page=page, per_page=20)
        employees = Employee.query.filter_by(active=True).all()

    elif report_type == 'payroll':
        # Example payroll: just employees with salary (you can expand later)
        query = Employee.query.filter(Employee.salary != None)
        data = query.paginate(page=page, per_page=20)
        employees = None

    else:
        flash("Invalid report type", "error")
        return redirect(url_for('hr_admin.reports'))

    return render_template(
        'hr/admin/reports/reports.html',
        data=data,
        report_type=report_type,
        start_date=start_date or '',
        end_date=end_date or '',
        employees=employees if report_type in ['attendance', 'leaves'] else []
    )



@hr_admin_bp.route('/attendance-report', methods=['GET'])
@login_required
@admin_required
def attendance_report():
    # ------------------------------
    # Get filter params
    # ------------------------------
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    department_id = request.args.get('department_id')

    if not start_date:
        start_date = date.today().replace(day=1)
    else:
        start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
    
    if not end_date:
        end_date = date.today()
    else:
        end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

    # ------------------------------
    # Fetch Employees (with eager loading)
    # ------------------------------
    employees_query = Employee.query.filter(Employee.status == "Active", Employee.archived == False)
    if department_id:
        employees_query = employees_query.filter(Employee.department_id == department_id)
    employees = employees_query.all()
    total_employees = len(employees)
    departments = Department.query.all()
    total_days = max(1, (end_date - start_date).days + 1)  # Prevent division by zero

    if total_employees == 0:
        return render_template(
            "hr/admin/reports/attendance_reports.html",
            report_data=[], department_summary=[], top_absentees=[],
            top_latecomers=[], perfect_attendees=[], best_dept=None, worst_dept=None,
            total_employees=0, total_hours_worked=0, avg_attendance_rate=0,
            avg_punctuality_rate=0, start_date=start_date, end_date=end_date,
            department_id=department_id, departments=departments,
            current_date=date.today().strftime("%B %d, %Y"), current_user=current_user
        )

    # ------------------------------
    # Efficient Attendance Fetch (1 Query)
    # ------------------------------
    all_attendance = Attendance.query.filter(
        Attendance.date >= start_date,
        Attendance.date <= end_date,
        Attendance.employee_id.in_([emp.id for emp in employees])
    ).all()

    emp_att_map = defaultdict(list)
    for att in all_attendance:
        emp_att_map[att.employee_id].append(att)

    # ------------------------------
    # Process Data & Insights
    # ------------------------------
    report_data = []
    top_absentees = []
    top_latecomers = []
    perfect_attendees = []
    total_hours_worked = 0
    total_present_days = 0

    for emp in employees:
        atts = emp_att_map.get(emp.id, [])
        
        # Count present/absent/late with safe status check
        days_present = sum(1 for a in atts if a.status and a.status.lower() in ["present", "late"])
        days_absent = sum(1 for a in atts if a.status and a.status.lower() == "absent")
        late_count = sum(1 for a in atts if a.status and a.status.lower() == "late")
        
        # ✅ FIX: Safe hours calculation with fallback
        hours_worked = sum((a.working_hours or 0) for a in atts)
        
        # If working_hours is still 0 but we have time_in/time_out, calculate fallback
        if hours_worked == 0 and atts:
            for a in atts:
                if a.time_in and a.time_out:
                    # Fallback: calculate from time_in/time_out (8hr day assumption)
                    from datetime import datetime as dt
                    start = dt.combine(a.date, a.time_in)
                    end = dt.combine(a.date, a.time_out)
                    diff_hours = (end - start).total_seconds() / 3600
                    hours_worked += max(0, diff_hours - 1) if diff_hours > 4 else diff_hours  # 1hr break

        punctuality_rate = round(((days_present - late_count) / days_present) * 100, 2) if days_present > 0 else 100

        total_hours_worked += hours_worked
        total_present_days += days_present

        record = {
            "employee_id": emp.id,
            "employee_name": emp.get_full_name(),
            "department_name": emp.department.name if emp.department else "N/A",
            "department_id": emp.department_id,
            "days_present": days_present,
            "days_absent": days_absent,
            "late_count": late_count,
            "total_hours": round(hours_worked, 2),
            "punctuality_rate": punctuality_rate
        }
        report_data.append(record)

        # Categorize for insights
        if days_absent > 0:
            top_absentees.append(record)
        if late_count > 0:
            top_latecomers.append(record)
        if days_absent == 0 and late_count == 0 and days_present > 0:
            perfect_attendees.append(record)

    # Sort insights
    top_absentees.sort(key=lambda x: x['days_absent'], reverse=True)
    top_latecomers.sort(key=lambda x: x['late_count'], reverse=True)

    # Overall Stats
    total_possible_days = total_employees * total_days
    avg_attendance_rate = round((total_present_days / total_possible_days) * 100, 2) if total_possible_days > 0 else 0
    avg_punctuality_rate = round(sum(r['punctuality_rate'] for r in report_data) / len(report_data), 2) if report_data else 0

    # Department Summary
    department_summary = []
    dept_metrics = {}
    
    for dept in departments:
        dept_emps = [r for r in report_data if r['department_id'] == dept.id]
        if not dept_emps:
            continue

        dept_att_days = sum(r['days_present'] for r in dept_emps)
        dept_hours = sum(r['total_hours'] for r in dept_emps)
        num_days = total_days * len(dept_emps)
        avg_att = round((dept_att_days / num_days) * 100, 2) if num_days > 0 else 0
        avg_hours = round(dept_hours / len(dept_emps), 2)

        department_summary.append({
            "name": dept.name, 
            "avg_attendance": avg_att, 
            "avg_hours": avg_hours,
            "employee_count": len(dept_emps)
        })
        dept_metrics[dept.id] = avg_att

    best_dept = None
    worst_dept = None
    if dept_metrics:
        best_dept_id = max(dept_metrics, key=dept_metrics.get)
        worst_dept_id = min(dept_metrics, key=dept_metrics.get)
        best_dept = next((d.name for d in departments if d.id == best_dept_id), None)
        worst_dept = next((d.name for d in departments if d.id == worst_dept_id), None)

    return render_template(
        "hr/admin/reports/attendance_reports.html",
        report_data=report_data,
        department_summary=department_summary,
        top_absentees=top_absentees[:5],
        top_latecomers=top_latecomers[:5],
        perfect_attendees=perfect_attendees,
        best_dept=best_dept,
        worst_dept=worst_dept,
        total_employees=total_employees,
        total_hours_worked=round(total_hours_worked, 2),
        avg_attendance_rate=avg_attendance_rate,
        avg_punctuality_rate=avg_punctuality_rate,
        start_date=start_date,
        end_date=end_date,
        department_id=int(department_id) if department_id else None,
        departments=departments,
        current_date=date.today().strftime("%B %d, %Y"),
        current_user=current_user
    )




@hr_admin_bp.route('/attendance/reports/word')
@login_required
@admin_required
def attendance_report_word():
    from datetime import datetime as dt, time as dt_time
    
    # -----------------------------
    # Filters from GET
    # -----------------------------
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    department_id = request.args.get('department_id')

    if start_date_str:
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
    else:
        start_date = date.today().replace(day=1)

    if end_date_str:
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
    else:
        end_date = date.today()

    total_days = max(1, (end_date - start_date).days + 1)

    # -----------------------------
    # Fetch employees
    # -----------------------------
    employees = Employee.query.filter(
        Employee.archived == False, 
        Employee.status == "Active"
    )
    if department_id:
        employees = employees.filter(Employee.department_id == department_id)
    employees = employees.all()

    # Pre-fetch all attendance in date range for efficiency
    emp_ids = [emp.id for emp in employees]
    all_att = Attendance.query.filter(
        Attendance.employee_id.in_(emp_ids),
        Attendance.date >= start_date,
        Attendance.date <= end_date
    ).all() if emp_ids else []
    
    att_map = defaultdict(list)
    for a in all_att:
        att_map[a.employee_id].append(a)

    # -----------------------------
    # ✅ HELPER: Calculate hours with fallback logic
    # -----------------------------
    def calculate_hours_safe(attendance):
        """
        Returns working hours with fallback calculation.
        Mirrors Attendance.calculate_working_hours() logic.
        """
        # If working_hours is already populated, use it
        if attendance.working_hours and attendance.working_hours > 0:
            return attendance.working_hours
        
        # Fallback: calculate from time_in/time_out
        if not attendance.time_in or not attendance.time_out:
            return 0.0
        
        # Get shift for this attendance (with fallback to default 8AM-5PM)
        shift = None
        if hasattr(attendance, 'get_shift'):
            shift = attendance.get_shift()
        
        # Default shift if none found: 8:00 AM to 5:00 PM (1hr break = 8hrs)
        work_start = dt_time(8, 0)
        work_end = dt_time(17, 0)
        
        if shift:
            work_start = shift.start_time
            work_end = shift.end_time
        
        # Combine date with times
        shift_start = dt.combine(attendance.date, work_start)
        shift_end = dt.combine(attendance.date, work_end)
        actual_in = dt.combine(attendance.date, attendance.time_in)
        actual_out = dt.combine(attendance.date, attendance.time_out)
        
        # Calculate overlap between actual time and shift
        start = max(actual_in, shift_start)
        end = min(actual_out, shift_end)
        
        if end <= start:
            return 0.0
        
        total_hours = (end - start).total_seconds() / 3600
        
        # Subtract 1 hour break if worked more than 4 hours (matches your model)
        if total_hours > 4:
            total_hours -= 1
        
        return round(max(0, total_hours), 2)

    # -----------------------------
    # Create Word Document
    # -----------------------------
    doc = Document()

    # Header
    header = doc.add_paragraph()
    header.alignment = 1  # center
    header.add_run("MUNICIPALITY OF NORZAGARAY\n").bold = True
    header.add_run("Attendance Report\n").bold = True
    header.add_run(f"Period: {start_date.strftime('%B %d, %Y')} to {end_date.strftime('%B %d, %Y')}\n").italic = True
    
    # ✅ Safe user name call
    user_name = 'HR Admin'
    if current_user and hasattr(current_user, 'get_full_name') and current_user.get_full_name:
        try:
            user_name = current_user.get_full_name()
        except:
            user_name = 'HR Admin'
    
    header.add_run(f"Generated: {date.today().strftime('%B %d, %Y')} by {user_name}").italic = True

    # Table header
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Employee Name"
    hdr_cells[1].text = "Department"
    hdr_cells[2].text = "Present"
    hdr_cells[3].text = "Absent"
    hdr_cells[4].text = "Late"
    hdr_cells[5].text = "Hours Worked"

    # Attendance data
    for emp in employees:
        emp_att = att_map.get(emp.id, [])
        
        days_present = sum(1 for a in emp_att if a.status and a.status.lower() in ["present", "late"])
        days_absent = sum(1 for a in emp_att if a.status and a.status.lower() == "absent")
        late_count = sum(1 for a in emp_att if a.status and a.status.lower() == "late")
        
        # ✅ FIX: Use safe calculation for hours
        total_hours = sum(calculate_hours_safe(a) for a in emp_att)

        row_cells = table.add_row().cells
        row_cells[0].text = emp.get_full_name()
        row_cells[1].text = emp.department.name if emp.department else "N/A"
        row_cells[2].text = str(days_present)
        row_cells[3].text = str(days_absent)
        row_cells[4].text = str(late_count)
        row_cells[5].text = f"{total_hours:.2f}"

    # -----------------------------
    # Summary Section
    # -----------------------------
    doc.add_paragraph('\n' + '='*60)
    doc.add_paragraph('SUMMARY', style='Heading 2')
    
    total_att_days = sum(
        sum(1 for a in att_map.get(emp.id, []) if a.status and a.status.lower() in ["present", "late"])
        for emp in employees
    )
    total_possible = total_days * len(employees) if employees else 1
    avg_att = round((total_att_days / total_possible) * 100, 2)
    
    # ✅ Use safe hours calculation for summary too
    total_hrs = sum(
        sum(calculate_hours_safe(a) for a in att_map.get(emp.id, []))
        for emp in employees
    )
    
    doc.add_paragraph(f"• Total Employees: {len(employees)}")
    doc.add_paragraph(f"• Reporting Period: {total_days} days")
    doc.add_paragraph(f"• Overall Attendance Rate: {avg_att}%")
    doc.add_paragraph(f"• Total Hours Worked: {total_hrs:.2f} hrs")
    doc.add_paragraph(f"• Avg Hours per Employee: {round(total_hrs/len(employees), 2) if employees else 0} hrs")

    # Department breakdown
    departments = Department.query.all()
    if departments:
        doc.add_paragraph('\nDEPARTMENT BREAKDOWN', style='Heading 2')
        for dept in departments:
            dept_emps = [e for e in employees if e.department_id == dept.id]
            if not dept_emps:
                continue
            
            dept_att = sum(
                sum(1 for a in att_map.get(emp.id, []) if a.status and a.status.lower() in ["present", "late"])
                for emp in dept_emps
            )
            dept_hrs = sum(
                sum(calculate_hours_safe(a) for a in att_map.get(emp.id, []))
                for emp in dept_emps
            )
            dept_days = total_days * len(dept_emps)
            dept_avg_att = round((dept_att / dept_days) * 100, 2) if dept_days > 0 else 0
            
            doc.add_paragraph(f"• {dept.name}: {dept_avg_att}% attendance, {dept_hrs:.1f} total hrs")

    # -----------------------------
    # Return as Word file
    # -----------------------------
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"Attendance_Report_{start_date}_{end_date}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )




# ------------------------------
# Leave Report - Main View
# ------------------------------
@hr_admin_bp.route("/hr_admin/leave_report")
@admin_required
@login_required
def leave_report():
    try:
        # --- Parse Filters (NO DEFAULTS) ---
        start_date_str = request.args.get('start_date')
        end_date_str = request.args.get('end_date')
        department_id = request.args.get('department_id', type=int)
        status_filter = request.args.get('status')
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 25, type=int)

        # Parse dates ONLY if explicitly provided
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date() if start_date_str else None
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date() if end_date_str else None

        if start_date and end_date and start_date > end_date:
            flash("Start date cannot be after end date.", "error")
            start_date, end_date = end_date, start_date
            start_date_str, end_date_str = str(start_date), str(end_date)

        # --- Build Base Query ---
        leave_query = Leave.query.join(Employee).join(Department, isouter=True).join(LeaveType, isouter=True).filter(
            Employee.archived == False
        )

        # --- Apply Conditional Filters ---
        if start_date and end_date:
            leave_query = leave_query.filter(Leave.start_date <= end_date, Leave.end_date >= start_date)
        elif start_date:
            leave_query = leave_query.filter(Leave.start_date >= start_date)
        elif end_date:
            leave_query = leave_query.filter(Leave.end_date <= end_date)

        if department_id:
            leave_query = leave_query.filter(Employee.department_id == department_id)
        if status_filter:
            leave_query = leave_query.filter(Leave.status == status_filter)

        # --- Pagination ---
        pagination = leave_query.order_by(Leave.start_date.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        leave_data = pagination.items

        # --- Insights & Dept Summary ---
        insights = _calculate_leave_insights(leave_query)
        dept_summary = _get_department_summary(leave_query)
        departments = Department.query.order_by(Department.name).all()

        return render_template(
            "hr/admin/reports/leave_reports.html",
            leave_data=leave_data,
            pagination=pagination,
            start_date=start_date,
            end_date=end_date,
            departments=departments,
            department_id=department_id,
            status_filter=status_filter,
            per_page=per_page,
            **insights,
            dept_summary=dept_summary,
            current_filters={
                'start_date': start_date_str,
                'end_date': end_date_str,
                'department_id': department_id,
                'status': status_filter
            }
        )

    except Exception as e:
        logger.error(f"Leave report error: {str(e)}", exc_info=True)
        flash("An error occurred while loading the report.", "error")
        return redirect(url_for('hr_admin_bp.dashboard'))


# ------------------------------
# Leave Report - Word Export
# ------------------------------
@hr_admin_bp.route("/leave-report/word")
@login_required
@admin_required
def leave_report_word():
    try:
        start_date_str = request.args.get("start_date")
        end_date_str = request.args.get("end_date")
        department_id = request.args.get("department_id", type=int)
        status_filter = request.args.get("status")

        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date() if start_date_str else None
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date() if end_date_str else None

        query = Leave.query.join(Employee).join(Department, isouter=True).join(LeaveType, isouter=True).filter(
            Employee.archived == False
        )

        # Apply same conditional filters as main view
        if start_date and end_date:
            query = query.filter(Leave.start_date <= end_date, Leave.end_date >= start_date)
        elif start_date:
            query = query.filter(Leave.start_date >= start_date)
        elif end_date:
            query = query.filter(Leave.end_date <= end_date)

        if department_id:
            query = query.filter(Employee.department_id == department_id)
        if status_filter:
            query = query.filter(Leave.status == status_filter)

        all_leaves = query.order_by(Leave.start_date.asc(), Employee.last_name.asc()).all()
        doc = _generate_leave_report_doc(all_leaves, start_date_str, end_date_str, department_id, status_filter)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        filename = f"Leave_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        logger.error(f"Word export error: {str(e)}", exc_info=True)
        flash("Failed to generate report. Please try again.", "error")
        return redirect(url_for('hr_admin_bp.leave_report'))


# ------------------------------
# Helper: Calculate Insights
# ------------------------------
def _calculate_leave_insights(query):
    """Efficiently compute report insights using SQLAlchemy aggregation."""
    total_leaves = query.count()
    
    if total_leaves == 0:
        return {
            "total_leaves": 0,
            "avg_days_per_leave": 0,
            "most_common_leave_type": "N/A",
            "approved_count": 0,
            "pending_count": 0,
            "rejected_count": 0,
            "total_days_requested": 0
        }

    stats = query.with_entities(
        func.sum(Leave.days_requested).label('total_days'),
        func.avg(Leave.days_requested).label('avg_days')
    ).first()
    
    status_counts = dict(query.with_entities(
        Leave.status, func.count(Leave.id)
    ).group_by(Leave.status).all())
    
    leave_types = [t[0] for t in query.with_entities(LeaveType.name).filter(LeaveType.name.isnot(None)).all()]
    most_common = Counter(leave_types).most_common(1)

    return {
        "total_leaves": total_leaves,
        "avg_days_per_leave": round(stats.avg_days or 0, 2),
        "most_common_leave_type": most_common[0][0] if most_common else "N/A",
        "approved_count": status_counts.get('Approved', 0),
        "pending_count": status_counts.get('Pending', 0),
        "rejected_count": status_counts.get('Rejected', 0),
        "total_days_requested": stats.total_days or 0
    }


# ------------------------------
# Helper: Department Summary
# ------------------------------
def _get_department_summary(query):
    """Generate department-wise breakdown."""
    from collections import defaultdict
    dept_summary = defaultdict(lambda: {"total": 0, "total_days": 0})
    
    for leave in query.with_entities(Leave, Department.name, Employee.department_id).all():
        dept_name = leave[1] or "Unassigned"
        dept_summary[dept_name]["total"] += 1
        dept_summary[dept_name]["total_days"] += leave[0].days_requested

    return {
        dept: {
            "total": data["total"],
            "avg_days": round(data["total_days"] / data["total"], 2) if data["total"] > 0 else 0,
            "total_days": data["total_days"]
        }
        for dept, data in dept_summary.items()
    }


# ------------------------------
# Helper: Generate Word Document
# ------------------------------
def _generate_leave_report_doc(leaves, start_date_str, end_date_str, dept_id, status_filter):
    """Create a professionally formatted Word document."""
    doc = Document()
    
    title = doc.add_heading("Municipality of Norzagaray", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Human Resources Department")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph()
    
    doc.add_heading("Leave Activity Report", level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"Period: {start_date_str or 'All Dates'} to {end_date_str or 'All Dates'}\n")
    if dept_id:
        dept = Department.query.get(dept_id)
        meta.add_run(f"Department: {dept.name if dept else 'N/A'}\n")
    if status_filter:
        meta.add_run(f"Status Filter: {status_filter.title()}\n")
    meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    headers = ["Employee", "Department", "Leave Type", "Start", "End", "Days", "Status", "Reason"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for leave in leaves:
        row = table.add_row().cells
        row[0].text = leave.employee.get_full_name()
        row[1].text = leave.employee.department.name if leave.employee.department else "N/A"
        row[2].text = leave.leave_type.name if leave.leave_type else "N/A"
        row[3].text = leave.start_date.strftime("%Y-%m-%d")
        row[4].text = leave.end_date.strftime("%Y-%m-%d")
        row[5].text = str(leave.days_requested)
        row[6].text = leave.status.title()
        reason = (leave.reason or "")[:100] + ("..." if len(leave.reason or "") > 100 else "")
        row[7].text = reason

    doc.add_page_break()
    doc.add_heading("Summary Insights", level=2)
    
    total = len(leaves)
    avg_days = round(sum(l.days_requested for l in leaves) / total, 2) if total else 0
    
    insights_table = doc.add_table(rows=4, cols=2)
    insights_table.style = 'Light Shading Accent 1'
    insight_data = [
        ("Total Leave Requests", str(total)),
        ("Total Days Requested", str(sum(l.days_requested for l in leaves))),
        ("Average Days per Request", str(avg_days)),
        ("Approved / Pending / Rejected", f"{sum(1 for l in leaves if l.status=='Approved')} / {sum(1 for l in leaves if l.status=='Pending')} / {sum(1 for l in leaves if l.status=='Rejected')}")
    ]
    for i, (label, value) in enumerate(insight_data):
        insights_table.rows[i].cells[0].text = label
        insights_table.rows[i].cells[1].text = value

    if leaves:
        doc.add_paragraph()
        doc.add_heading("Department Breakdown", level=2)
        dept_data = _get_department_summary(Leave.query.filter(Leave.id.in_([l.id for l in leaves])))
        
        dept_table = doc.add_table(rows=1, cols=3)
        dept_table.style = 'Light Grid'
        for cell in dept_table.rows[0].cells:
            cell.text = "Department" if cell == dept_table.rows[0].cells[0] else "Total Leaves" if cell == dept_table.rows[0].cells[1] else "Avg Days"
            cell.paragraphs[0].runs[0].bold = True
            
        for dept, stats in sorted(dept_data.items()):
            row = dept_table.add_row().cells
            row[0].text = dept
            row[1].text = str(stats['total'])
            row[2].text = str(stats['avg_days'])

    return doc