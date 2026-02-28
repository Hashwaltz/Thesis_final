from flask import Blueprint, render_template, request, session, current_app, redirect, url_for, flash, jsonify, Response, send_file, make_response
from flask_login import login_required, current_user
from sqlalchemy.exc import IntegrityError
from datetime import datetime
from main_app.models.user import User
from main_app.models.hr_models import Employee, Attendance, Leave, Department, Position, LeaveType, EmploymentType, LeaveCredit
from main_app.forms import EmployeeForm, AttendanceForm, LeaveForm, DepartmentForm
from main_app.utils import admin_required, generate_employee_id, get_attendance_summary, get_current_month_range, load_excel_to_df, unlock_xlsx
from main_app.extensions import db, mail
from datetime import timedelta, datetime, date, time
from sqlalchemy.orm import joinedload
from collections import defaultdict
from main_app.functions import parse_date
import os
from collections import Counter
import csv
from docx.shared import Pt, RGBColor
import shutil
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from flask import send_file
import io
from io import BytesIO
import pdfkit
from werkzeug.utils import secure_filename
from docx import Document
import pandas as pd
import re
import numpy as np
import uuid
from sqlalchemy import func, and_, case, cast, Date
import json
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.pagesizes import LETTER
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_mail import Message

pdfmetrics.registerFont(UnicodeCIDFont('HYSMyeongJo-Medium'))  # For Filipino/Unicode chars
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../"))
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")

hr_admin_bp = Blueprint(
    'hr_admin',
    __name__,
    template_folder=TEMPLATE_DIR,
    static_url_path='/hr/static'
)


@hr_admin_bp.route('/dashboard')
@login_required
@admin_required
def hr_dashboard():
    today = date.today()

    # --- Basic Stats ---
    total_employees = Employee.query.count()
    # Count only Active employees using status string
    active_employees = Employee.query.filter_by(status="Active").count()
    total_departments = Department.query.count()

    # --- Recent records ---
    recent_employees = Employee.query.order_by(Employee.created_at.desc()).limit(5).all()
    recent_leaves = Leave.query.order_by(Leave.created_at.desc()).limit(5).all()

    # --- Employee Growth (Monthly Count) ---
    growth_labels = []
    growth_counts = []
    for i in range(6, 0, -1):
        month = today.replace(day=1) - timedelta(days=30*(i-1))
        month_start = month.replace(day=1)
        month_end = (month_start + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        count = Employee.query.filter(Employee.date_hired.between(month_start, month_end)).count()
        growth_labels.append(month_start.strftime("%b"))
        growth_counts.append(count)

    # --- Department Distribution ---
    dept_data = db.session.query(
        Department.name, db.func.count(Employee.id)
    ).join(Employee, Employee.department_id == Department.id)\
     .group_by(Department.name).all()
    dept_labels = [d[0] for d in dept_data]
    dept_counts = [d[1] for d in dept_data]

    # --- Employees per Barangay ---
    barangay_data = db.session.query(
        Employee.barangay, db.func.count(Employee.id)
    ).group_by(Employee.barangay).all()
    barangay_labels = [b[0] or "N/A" for b in barangay_data]
    barangay_counts = [b[1] for b in barangay_data]

    # --- Attendance Overview (Past 7 days) ---
    attendance_labels = []
    attendance_counts = []
    for i in range(7):
        day = today - timedelta(days=6-i)
        records = Attendance.query.filter_by(date=day).all()
        total = len(records)
        present = len([r for r in records if r.status == "Present"])
        attendance_percentage = round((present / total * 100) if total else 0, 2)
        attendance_labels.append(day.strftime("%a"))
        attendance_counts.append(attendance_percentage)

    # --- Leave Requests ---
    leave_data = db.session.query(
        Leave.status, db.func.count(Leave.id)
    ).group_by(Leave.status).all()
    leave_labels = [l[0] for l in leave_data]
    leave_counts = [l[1] for l in leave_data]

    return render_template(
        'hr/admin/admin_dashboard.html',
        total_employees=total_employees,
        active_employees=active_employees,
        total_departments=total_departments,
        recent_employees=recent_employees,
        recent_leaves=recent_leaves,
        growth_labels=growth_labels,
        growth_counts=growth_counts,
        dept_labels=dept_labels,
        dept_counts=dept_counts,
        barangay_labels=barangay_labels,
        barangay_counts=barangay_counts,
        attendance_labels=attendance_labels,
        attendance_counts=attendance_counts,
        leave_labels=leave_labels,
        leave_counts=leave_counts,
        user=current_user
    )

@hr_admin_bp.route('/employees')
@login_required
@admin_required
def view_employees():
    search = request.args.get('search', '')
    department_id = request.args.get('department_id', '')
    employment_type_id = request.args.get('employment_type_id', '')
    page = request.args.get('page', 1, type=int)

    # Base query with joins, excluding archived (archived=True will be hidden)
    query = Employee.query.options(
        joinedload(Employee.department),
        joinedload(Employee.position),
        joinedload(Employee.employment_type)
    ).filter(Employee.archived.isnot(True))  # ✅ includes False and NULL

    # Apply search filter
    if search:
        query = query.filter(
            Employee.first_name.ilike(f"%{search}%") |
            Employee.last_name.ilike(f"%{search}%") |
            Employee.email.ilike(f"%{search}%")
        )

    # Apply department filter
    if department_id:
        query = query.filter(Employee.department_id == int(department_id))

    # Apply employment type filter
    if employment_type_id:
        query = query.filter(Employee.employment_type_id == int(employment_type_id))

    # Order employees alphabetically
    query = query.order_by(Employee.last_name.asc(), Employee.first_name.asc())

    # Paginate
    employees = query.paginate(page=page, per_page=10)

    # Fetch all filter dropdown data
    departments = Department.query.order_by(Department.name.asc()).all()
    employment_types = EmploymentType.query.order_by(EmploymentType.name.asc()).all()
    positions = Position.query.order_by(Position.name.asc()).all()

    return render_template(
        'hr/admin/admin_view_employees.html',
        employees=employees,
        search=search,
        positions=positions,
        departments=departments,
        employment_types=employment_types,
        selected_department=department_id,
        selected_employment_type=employment_type_id
    )


@hr_admin_bp.route('/employees/add', methods=['POST'])
@login_required
@admin_required
def add_employee():
    try:
        # --- 1. Get form data ---
        department_id = request.form['department_id']
        new_employee_id = generate_employee_id(department_id)

        # Parse dates safely
        date_hired = parse_date(request.form['date_hired'], "Date Hired")
        date_of_birth = parse_date(request.form['date_of_birth'], "Date of Birth")
        if not date_hired or not date_of_birth:
            flash("Invalid date format!", "danger")
            return redirect(url_for('hr_admin.view_employees'))

        # Parse salary
        salary_str = request.form.get('salary', '').strip()
        try:
            salary = float(salary_str) if salary_str else 0.0
        except ValueError:
            flash("Invalid salary value!", "danger")
            return redirect(url_for('hr_admin.view_employees'))

        # Address
        street = request.form.get('street_address', '').strip()
        barangay = request.form.get('barangay', '').strip()
        municipality = request.form.get('municipality', '').strip()
        province = request.form.get('province', '').strip()
        postal_code = request.form.get('postal_code', '').strip()

        # --- 2. Create User ---
        default_password = "password123"
        user = User(
            email=request.form['email'],
            first_name=request.form['first_name'],
            last_name=request.form['last_name'],
            role="employee",
            password=default_password
        )
        db.session.add(user)
        db.session.flush()  # get user.id

        # --- 3. Create Employee ---
        employment_type_id = int(request.form['employment_type_id'])
        employee = Employee(
            employee_id=new_employee_id,
            user_id=user.id,
            first_name=request.form['first_name'],
            last_name=request.form['last_name'],
            middle_name=request.form.get('middle_name'),
            email=request.form['email'],
            phone=request.form['phone'],
            street_address=street,
            barangay=barangay,
            municipality=municipality,
            province=province,
            postal_code=postal_code,
            department_id=department_id,
            position_id=request.form['position_id'],
            employment_type_id=employment_type_id,
            salary=salary,
            date_hired=date_hired,
            date_of_birth=date_of_birth,
            gender=request.form['gender'],
            marital_status=request.form['marital_status'],
            emergency_contact=request.form['emergency_contact'],
            status='Active'
        )
        db.session.add(employee)
        db.session.flush()  # get employee.id for leave credits

        # --- 4. Initialize Leave Credits ONLY for Regular (1) & Casual (3) ---
        if employment_type_id in [1, 3]:
            # Vacation Leave (leave_type_id = 1)
            vacation_credit = LeaveCredit(
                employee_id=employee.id,
                leave_type_id=1,
                total_credits=15.0,
                used_credits=0
            )
            db.session.add(vacation_credit)

            # Sick Leave (leave_type_id = 2)
            sick_credit = LeaveCredit(
                employee_id=employee.id,
                leave_type_id=2,
                total_credits=15.0,
                used_credits=0
            )
            db.session.add(sick_credit)

        # --- 5. Commit everything together ---
        db.session.commit()

        # --- 6. Send Gmail notification ---
        try:
            msg = Message(
                subject="Your govHRPay Account Details",
                sender=("GovHRPay Admin", "natanielashleyrodelas@gmail.com"),  # system email
                recipients=[user.email]
            )
            msg.body = f"""
Hello {user.first_name} {user.last_name},

Your govHRPay account has been created successfully!

Login credentials:
Email: {user.email}
Password: {default_password}

Please log in at: http://127.0.0.1:5000/hr/auth/login

For security, you should change your password after first login.

Thank you,
GovHRPay Admin
"""
            mail.send(msg)
        except Exception as mail_err:
            current_app.logger.error(f"Failed to send account email to {user.email}: {mail_err}")
            flash("Employee created, but failed to send email notification.", "warning")
        else:
            flash("Employee and user account created successfully! Email sent with login details.", "success")

        return redirect(url_for('hr_admin.view_employees'))

    except IntegrityError as e:
        db.session.rollback()
        flash(f"Error: Employee or User already exists! ({str(e)})", "danger")
        return redirect(url_for('hr_admin.view_employees'))

    except Exception as e:
        db.session.rollback()
        flash(f"Unexpected error: {str(e)}", "danger")
        return redirect(url_for('hr_admin.view_employees'))


@hr_admin_bp.route("/generate_moa_all/<int:employment_type_id>")
@login_required
@admin_required
def generate_moa_all(employment_type_id):
    from flask import send_file, flash, redirect, url_for
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    from datetime import date

    # ===== FETCH DATA =====
    if employment_type_id == 0:  # 0 = all types
        etypes = EmploymentType.query.order_by(EmploymentType.name).all()
        employees_by_type = {etype: Employee.query.filter_by(employment_type_id=etype.id).all() for etype in etypes}
    else:
        etype = EmploymentType.query.get_or_404(employment_type_id)
        employees_by_type = {etype: Employee.query.filter_by(employment_type_id=etype.id).all()}

    if not any(employees_by_type.values()):
        flash("No employees found for the selected type(s).", "warning")
        return redirect(url_for("hr_admin.view_employees"))

    # ===== CREATE WORKBOOK =====
    wb = Workbook()
    bold_center = Font(bold=True)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )
    gray_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

    # ===== LOOP THROUGH TYPES =====
    for i, (etype, employees) in enumerate(employees_by_type.items()):
        if i == 0:
            ws = wb.active
            ws.title = etype.name[:28]  # max sheet name length
        else:
            ws = wb.create_sheet(title=etype.name[:28])

        # ===== MAIN HEADER =====
        ws.merge_cells("A1:K1")
        ws["A1"] = f"LIST OF {etype.name.upper()} PERSONNEL"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A1"].alignment = Alignment(horizontal="center")

        ws.merge_cells("A2:K2")
        ws["A2"] = f"(As of {date.today().strftime('%B %d, %Y')})"
        ws["A2"].font = Font(italic=True)
        ws["A2"].alignment = Alignment(horizontal="center")

        # ===== AGENCY INFO =====
        ws.append([])
        ws.append(["Agency name:", "LGU-NORZAGARAY, BULACAN"])
        ws.append(["Regional Office No:", "3"])
        ws.append([])

        start_row = ws.max_row + 1

        # ===== TWO-LEVEL HEADER =====
        ws.merge_cells(f"A{start_row}:A{start_row+1}")
        ws.merge_cells(f"B{start_row}:D{start_row}")
        ws.merge_cells(f"E{start_row}:E{start_row+1}")
        ws.merge_cells(f"F{start_row}:F{start_row+1}")
        ws.merge_cells(f"G{start_row}:G{start_row+1}")
        ws.merge_cells(f"H{start_row}:H{start_row+1}")
        ws.merge_cells(f"I{start_row}:I{start_row+1}")
        ws.merge_cells(f"J{start_row}:K{start_row}")

        ws[f"A{start_row}"] = "NO."
        ws[f"B{start_row}"] = "Name of Personnel"
        ws[f"E{start_row}"] = "DATE OF BIRTH\n(MM/DD/YYYY)"
        ws[f"F{start_row}"] = "SEX\n(pls. select)"
        ws[f"G{start_row}"] = "Level of CS Eligibility\n(pls. select)"
        ws[f"H{start_row}"] = "WORK STATUS\n(pls. select)"
        ws[f"I{start_row}"] = f"No. of Years of Service as {etype.name} personnel"
        ws[f"J{start_row}"] = "NATURE OF WORK"

        ws[f"B{start_row+1}"] = "SURNAME"
        ws[f"C{start_row+1}"] = "FIRST NAME/\nEXTENSION NAME"
        ws[f"D{start_row+1}"] = "MIDDLE INITIAL"
        ws[f"J{start_row+1}"] = "Pls select"
        ws[f"K{start_row+1}"] = "Please specify"

        # ===== STYLE HEADER =====
        for row in ws.iter_rows(min_row=start_row, max_row=start_row+1, min_col=1, max_col=11):
            for cell in row:
                cell.font = bold_center
                cell.alignment = center_align
                cell.border = thin_border
                cell.fill = gray_fill

        # ===== TABLE BODY =====
        for idx, emp in enumerate(employees, start=1):
            ws.append([
                idx,
                emp.last_name or "",
                emp.first_name or "",
                emp.middle_name or "",
                emp.date_of_birth.strftime("%m/%d/%Y") if emp.date_of_birth else "",
                emp.gender or "",
                getattr(emp, "cs_eligibility", "No eligibility"),
                emp.status or "Active",
                emp.get_working_duration() if hasattr(emp, "get_working_duration") else "",
                emp.position.name if emp.position else "",
                "",  # "Please specify"
            ])

        # ===== AUTO WIDTH =====
        for i_col, col in enumerate(ws.columns, start=1):
            max_length = max((len(str(cell.value)) for cell in col if cell.value), default=0)
            ws.column_dimensions[get_column_letter(i_col)].width = max_length + 2

    # ===== SAVE & RETURN =====
    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    filename = f"MOA_List_{date.today().strftime('%Y%m%d')}.xlsx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@hr_admin_bp.route('/employees/export')
@login_required
@admin_required
def export_employees_excel():
    
    employees = Employee.query.filter_by(status='active').all()

    wb = Workbook()
    ws = wb.active
    ws.title = "Employees"

    headers = ['Employee ID', 'First Name', 'Last Name', 'Email', 'Department', 'Status']
    ws.append(headers)

    for emp in employees:
        ws.append([
            emp.employee_id,
            emp.first_name,
            emp.last_name,
            emp.email or '',
            emp.department.name if emp.department else '',
            'Active' if emp.active else 'Inactive'
        ])

    # 🔹 Auto-adjust column width based on max content length
    for col in ws.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[col_letter].width = adjusted_width

    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="employees_report.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

@hr_admin_bp.route('/employees/<int:employee_id>/edit', methods=['GET', 'POST'])
@admin_required
@login_required
def edit_employee(employee_id):
    employee = Employee.query.get_or_404(employee_id)

    if request.method == 'GET':
        # Return JSON for modal population
        return jsonify({
            "first_name": employee.first_name,
            "middle_name": employee.middle_name,
            "last_name": employee.last_name,
            "email": employee.email,
            "phone": employee.phone,
            "gender": employee.gender,
            "marital_status": employee.marital_status,
            "emergency_contact": employee.emergency_contact,
            "emergency_phone": employee.emergency_phone,
            "department_id": employee.department_id,
            "position_id": employee.position_id,
            "employment_type_id": employee.employment_type_id,
            "salary": str(employee.salary) if employee.salary else "",
            "date_of_birth": employee.date_of_birth.isoformat() if employee.date_of_birth else "",
            "date_hired": employee.date_hired.isoformat() if employee.date_hired else "",
            "status": employee.status,  
            "street_address": employee.street_address,
            "barangay": employee.barangay,
            "municipality": employee.municipality,
            "province": employee.province,
            "postal_code": employee.postal_code
        })

    if request.method == 'POST':
        try:
            # Personal info
            new_email = request.form.get("email")
            employee.first_name = request.form.get("first_name")
            employee.middle_name = request.form.get("middle_name")
            employee.last_name = request.form.get("last_name")
            employee.email = new_email
            employee.phone = request.form.get("phone")
            employee.gender = request.form.get("gender")
            employee.marital_status = request.form.get("marital_status")
            employee.emergency_contact = request.form.get("emergency_contact")
            employee.emergency_phone = request.form.get("emergency_phone")

            # Address fields
            employee.street_address = request.form.get("street_address")
            employee.barangay = request.form.get("barangay")
            employee.municipality = request.form.get("municipality")
            employee.province = request.form.get("province")
            employee.postal_code = request.form.get("postal_code")

            # Employment info
            employee.department_id = int(request.form.get("department_id")) if request.form.get("department_id") else None
            employee.position_id = int(request.form.get("position_id")) if request.form.get("position_id") else None
            employee.employment_type_id = int(request.form.get("employment_type_id")) if request.form.get("employment_type_id") else None
            salary_val = request.form.get("salary")
            employee.salary = float(salary_val) if salary_val else None

            # Dates
            def parse_date(date_str):
                if not date_str:
                    return None
                try:
                    return datetime.strptime(date_str, "%Y-%m-%d").date()
                except ValueError:
                    return None

            employee.date_of_birth = parse_date(request.form.get("date_of_birth"))
            employee.date_hired = parse_date(request.form.get("date_hired"))

            # Status (string)
            employee.status = request.form.get("status")  # ✅ Updated

            # --- Update linked user email if exists ---
            if hasattr(employee, 'user') and employee.user:
                employee.user.email = new_email

            # Commit changes
            db.session.commit()
            return redirect(url_for('hr_admin.view_employees'))

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating employee {employee_id}: {e}")
            return redirect(url_for('hr_admin.view_employees'))


@hr_admin_bp.route('/employees/<int:employee_id>/service_record')
@login_required
@admin_required
def export_service_record(employee_id):
    employee = Employee.query.get_or_404(employee_id)

    # Create Word document
    doc = Document()

    # ==============================
    # HEADER SECTION
    # ==============================
    title = doc.add_paragraph("S E R V I C E   R E C O R D")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    para1 = doc.add_paragraph("Republic of the Philippines")
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para1.runs[0].font.size = Pt(11)

    para2 = doc.add_paragraph("NORZAGARAY, REGION 3")
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.runs[0].bold = True
    para2.runs[0].font.size = Pt(11)

    doc.add_paragraph("")  # spacing

    # ==============================
    # EMPLOYEE INFORMATION
    # ==============================
    doc.add_paragraph(f"Name : {employee.last_name.upper()}, {employee.first_name.upper()} {employee.middle_name or ''}")
    birth_date = employee.date_of_birth.strftime('%B %d, %Y') if employee.date_of_birth else ''
    doc.add_paragraph(f"Date and place of birth : {birth_date}")
    doc.add_paragraph("(If married woman, give full maiden name)")
    doc.add_paragraph("(Date herein should be checked from birth or baptismal certificate or some other reliable documents)")
    doc.add_paragraph("B.P. Number: __________     TIN #: __________")
    doc.add_paragraph("")  # spacing

    # ==============================
    # CERTIFICATION STATEMENT
    # ==============================
    cert_text = (
        "This is to certify that the employee named hereunder actually rendered services "
        "in this Office as shown by the service record below, each line of which is supported "
        "by appointment and other papers actually issued by this Office and approved by the "
        "authorities concerned."
    )
    cert_para = doc.add_paragraph(cert_text)
    cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")  # spacing

    # ==============================
    # SERVICE RECORD TABLE
    # ==============================
    headers = [
        "From", "To", "Designation Status (1)", "Annual Salary (2)",
        "Station / Place of Assignment", "Branch (3)", "Leave(s) w/out Pay Date", "Cause"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text

    # Add data row (current employment)
    row = table.add_row().cells
    row[0].text = employee.date_hired.strftime('%b %d, %Y') if employee.date_hired else ''
    row[1].text = "Present"
    row[2].text = employee.position.name if employee.position else ''
    row[3].text = f"{employee.salary or ''}"
    row[4].text = employee.department.name if employee.department else ''
    row[5].text = ""
    row[6].text = ""
    row[7].text = ""

    doc.add_paragraph("")  # spacing

    # ==============================
    # FOOTER SECTION
    # ==============================
    footer = (
        "Issued on compliance with Executive Order No. 54 dated August 10, 1954, and in accordance "
        "with Circular No. 68 dated August 10, 1954 of the System."
    )
    footer_para = doc.add_paragraph(footer)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("CERTIFIED CORRECT:")
    doc.add_paragraph("FERNANDO DG. CRUZ")
    doc.add_paragraph("Acting MHRMO")
    doc.add_paragraph("Page 1 of 1")
    doc.add_paragraph(date.today().strftime("%A, %B %d, %Y"))

    # Adjust font size
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

    # ==============================
    # SAVE TO BYTES AND RETURN
    # ==============================
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name=f"service_record_{employee.employee_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@hr_admin_bp.route("/employee/<int:employee_id>/generate-coe")
@login_required
def generate_coe(employee_id):

    import os
    import io
    from datetime import datetime
    from flask import current_app, send_file

    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Image,
        Table,
        TableStyle
    )

    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib import colors

    employee = Employee.query.get_or_404(employee_id)

    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=40,
        bottomMargin=60
    )

    styles = getSampleStyleSheet()

    # ===============================
    # Styles
    # ===============================

    center_style = ParagraphStyle(
        "Center",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=11
    )

    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        spaceAfter=20
    )

    right_style = ParagraphStyle(
        "Right",
        parent=styles["Normal"],
        alignment=TA_RIGHT,
        fontSize=10
    )

    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        leading=18,
        fontSize=11
    )

    # ===============================
    # Logo
    # ===============================

    logo_path = os.path.join(
        current_app.root_path,
        "static",
        "img",
        "garay.png"
    )

    if os.path.exists(logo_path):
        logo = Image(logo_path, width=70, height=70)
        header_table = Table([[logo]], colWidths=[450], hAlign="CENTER")
    else:
        header_table = Table([[""]], colWidths=[450])

    # ===============================
    # Header Text
    # ===============================

    header_text = """
    Republic of the Philippines<br/>
    Province of Bulacan<br/>
    MUNICIPALITY OF NORZAGARAY<br/>
    HUMAN RESOURCE MANAGEMENT OFFICE
    """

    separator = Table([[""]], colWidths=[450])
    separator.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.darkblue),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2)
    ]))

    # ===============================
    # Date
    # ===============================

    today_date = datetime.now().strftime("%B %d, %Y")
    date_paragraph = Paragraph(today_date, right_style)

    # ===============================
    # Gender Title Logic
    # ===============================

    gender = (employee.gender or "").lower()

    if gender == "male":
        title_prefix = "Mr."
    elif gender == "female":
        title_prefix = "Ms."
    else:
        title_prefix = "Mr./Ms."

    display_name = f"{title_prefix} {employee.last_name}"

    # ===============================
    # Model Data Extraction ⭐ IMPORTANT
    # ===============================

    department_name = employee.department.name if employee.department else "N/A"
    position_name = employee.position.name if employee.position else "N/A"
    employment_type = employee.employment_type.name if employee.employment_type else "N/A"

    hire_date = employee.date_hired.strftime("%B %d, %Y") if employee.date_hired else "N/A"

    end_date = "Present" if employee.status == "Active" else (employee.status or "N/A")

    working_duration = employee.get_working_duration()

    # ===============================
    # Certificate Body (Correct Version)
    # ===============================

    body_text = f"""
    <b>TO WHOM IT MAY CONCERN:</b><br/><br/>

    This is to certify that <b>{display_name}</b>,
    a <b>{position_name}</b> under the <b>{department_name}</b>,
    is employed as <b>{employment_type}</b> status in this office.

    <br/><br/>

    This employee has been working since <b>{hire_date}</b>
    up to <b>{end_date}</b> with a total working duration of
    <b>{working_duration}</b>.

    <br/><br/>

    This certification is issued upon the request of {display_name}
    for whatever legal purpose this may serve.
    """

    # ===============================
    # Signature Block
    # ===============================

    signature_block = Paragraph("""
    <br/><br/><br/>
    <b>FERNANDO DG. CRUZ</b><br/>
    Acting MHRMO
    """, right_style)

    # ===============================
    # Build PDF
    # ===============================

    story = [
        header_table,
        Paragraph(header_text, center_style),

        Spacer(1, 6),
        separator,
        Spacer(1, 20),

        date_paragraph,
        Spacer(1, 20),

        Paragraph("CERTIFICATION", title_style),

        Paragraph(body_text, body_style),

        signature_block
    ]

    doc.build(story)

    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"COE_{employee.employee_id}.pdf",
        mimetype="application/pdf"
    )

@hr_admin_bp.route("/employees/<int:employee_id>/archive", methods=["POST"])
@login_required
def archive_employee(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    employee.archived = True
    employee.archived_at = datetime.utcnow()

    db.session.commit()

    # Return JSON if AJAX, otherwise redirect to archived page
    if request.is_json or request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"success": True})
    
    return redirect(url_for("hr_admin.view_archived_employees"))

@hr_admin_bp.route("/employees/archived")
@login_required
def view_archived_employees():
    page = request.args.get("page", 1, type=int)
    search = request.args.get("search", "", type=str)
    department_id = request.args.get("department_id", type=int)
    employment_type_id = request.args.get("employment_type_id", type=int)

    query = Employee.query.filter_by(archived=True)

    if search:
        query = query.filter(
            (Employee.first_name.ilike(f"%{search}%")) |
            (Employee.last_name.ilike(f"%{search}%"))
        )
    if department_id:
        query = query.filter(Employee.department_id == department_id)
    if employment_type_id:
        query = query.filter(Employee.employment_type_id == employment_type_id)

    employees = query.order_by(Employee.archived_at.desc()).paginate(page=page, per_page=10)

    # Assuming you have department and employment type lists for filter dropdowns
    departments = Department.query.all()
    employment_types = EmploymentType.query.all()

    return render_template(
        "hr/admin/view_archives.html",
        employees=employees,
        departments=departments,
        employment_types=employment_types,
        search=search,
        selected_department=str(department_id) if department_id else "",
        selected_employment_type=str(employment_type_id) if employment_type_id else ""
    )


@hr_admin_bp.route('/employees/restore/<int:employee_id>', methods=['POST'])
@login_required
@admin_required
def restore_employee(employee_id):
    employee = Employee.query.get_or_404(employee_id)

    if not employee.archived:
        flash("Employee is not archived.", "warning")
        return redirect(url_for('hr_admin.view_archived_employees'))

    employee.archived = False
    db.session.commit()

    flash(f"Employee {employee.get_full_name()} has been restored.", "success")
    return redirect(url_for('hr_admin.view_archived_employees'))



@hr_admin_bp.route('/users', methods=['GET'])
@login_required
@admin_required
def view_users():
    # Get query params
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '').strip()
    role_filter = request.args.get('role', '').strip()
    status_filter = request.args.get('status', '').strip()  

    # Base query
    query = User.query

    # Apply search filter
    if search:
        query = query.filter(
            (User.first_name.ilike(f"%{search}%")) |
            (User.last_name.ilike(f"%{search}%")) |
            (User.email.ilike(f"%{search}%"))
        )

    # Apply role filter
    if role_filter:
        query = query.filter(User.role == role_filter)

    # ✅ Apply status filter
    if status_filter == "active":
        query = query.filter(User.active.is_(True))
    elif status_filter == "inactive":
        query = query.filter(User.active.is_(False))

    # Paginate results
    users = query.order_by(User.id.asc()).paginate(page=page, per_page=10)

    # Roles for dropdown
    roles = ['admin', 'employee', 'dept_head', 'officer']

    return render_template(
        'hr/admin/admin_view_users.html',
        users=users,
        roles=roles,
        search=search,
        role_filter=role_filter,
        status_filter=status_filter  
    )


@hr_admin_bp.route("/user/<int:user_id>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def edit_user(user_id):
    user = User.query.get_or_404(user_id)

    if request.method == "POST":
        try:
            user.email = request.form.get("email")
            user.first_name = request.form.get("first_name")
            user.last_name = request.form.get("last_name")
            user.role = request.form.get("role")
            user.active = request.form.get("status") == "1"

            db.session.commit()

            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "success", "message": "User updated successfully!"})

            flash("User updated successfully!", "success")
            return redirect(url_for("hr_admin.view_users"))

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating user {user_id}: {e}")
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "error", "message": "Error updating user."}), 500
            flash("Error updating user.", "error")

    # GET → JSON for modal
    if request.headers.get("Accept") == "application/json":
        return jsonify({
            "email": user.email,
            "first_name": user.first_name,
            "last_name": user.last_name,
            "role": user.role,
            "active": user.active
        })

    return redirect(url_for("hr_admin.view_users"))


@hr_admin_bp.route('/attendance')
@login_required
@admin_required
def attendance():
    page = request.args.get('page', 1, type=int)
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    employee_filter = request.args.get('employee', '').strip()
    department_filter = request.args.get('department', '').strip()
    status_filter = request.args.get('status', '').strip()  

    # Base query
    query = Attendance.query.join(Employee).join(Employee.department)

    # Date filters
    try:
        if start_date and not end_date:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            query = query.filter(Attendance.date == start_date_obj)

        elif end_date and not start_date:
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            query = query.filter(Attendance.date == end_date_obj)

        elif start_date and end_date:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()

            # Swap if user entered in reverse
            if end_date_obj < start_date_obj:
                start_date_obj, end_date_obj = end_date_obj, start_date_obj

            query = query.filter(
                and_(
                    Attendance.date >= start_date_obj,
                    Attendance.date <= end_date_obj
                )
            )
    except ValueError:
        pass

    # Employee and Department filters
    if employee_filter:
        query = query.filter(Attendance.employee_id == int(employee_filter))
    if department_filter:
        query = query.filter(Employee.department_id == int(department_filter))

    # ✅ Status filter
    if status_filter:
        query = query.filter(Attendance.status == status_filter)

    # Pagination
    attendances = query.order_by(Attendance.date.desc()).paginate(page=page, per_page=20, error_out=False)

    # Lists for dropdowns
    employees = Employee.query.filter_by(archived=False).all()
    departments = Department.query.order_by(Department.name.asc()).all()

    return render_template(
        'hr/admin/admin_view_attendance.html',
        attendances=attendances,
        employees=employees,
        departments=departments,
        start_date=start_date,
        end_date=end_date,
        employee_filter=employee_filter,
        department_filter=department_filter,
        status_filter=status_filter  
    )


# ----------------- CONFIG -----------------
ALLOWED_EXTENSIONS = {'xls', 'xlsx'}
UPLOAD_FOLDER = "uploads/attendance"

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



# ----------------- CONFIG -----------------
ALLOWED_EXTENSIONS = {'xls', 'xlsx'}
UPLOAD_FOLDER = "uploads/attendance"

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@hr_admin_bp.route('/add_attendance', methods=['GET', 'POST'])
@login_required
def add_attendance():
    preview_data = []
     
    employees = Employee.query.filter_by(status='Active').all()
    if request.method == 'POST' and 'file' in request.files:
        file = request.files.get("file")
        if not file or not allowed_file(file.filename):
            flash("Please upload a valid Excel file (.xls or .xlsx).", "danger")
            return redirect(request.url)

        # Save uploaded file
        filename = secure_filename(file.filename)
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        filepath = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}_{filename}")
        file.save(filepath)

        try:
            df = pd.read_excel(filepath, header=None)
            records = []
            current_id, current_name, current_dept = None, None, None
            attendance_date = None

            for _, row in df.iterrows():
                line = " ".join(str(x) for x in row if str(x) != "nan").strip()
                if not line or "tabling date" in line.lower():
                    continue

                if "Attendance date:" in line:
                    attendance_date = line.split(":")[-1].strip()
                    continue

                if "User ID" in line and "Name" in line:
                    uid_part = line.split("User ID:")[-1]
                    name_part = uid_part.split("Name:")
                    id_value = name_part[0].strip() if len(name_part) > 0 else None

                    if len(name_part) > 1:
                        name_dept_part = name_part[1].split("Department:")
                        name = name_dept_part[0].strip()
                        dept = name_dept_part[1].strip() if len(name_dept_part) > 1 else "Unknown"
                    else:
                        name, dept = "Unknown", "Unknown"

                    current_id, current_name, current_dept = id_value, name, dept
                    continue

                if ":" in line:
                    times = line.split()
                    time_in = times[0] if len(times) > 0 else None
                    time_out = times[1] if len(times) > 1 else None
                    day_value = attendance_date or datetime.now().date().isoformat()

                    emp_match = None
                    try:
                        emp_id_int = int(float(current_id))
                        emp_match = Employee.query.get(emp_id_int)
                    except:
                        emp_match = None

                    record_name = emp_match.get_full_name() if emp_match else current_name or "Unknown"
                    matched = True if emp_match else False

                    records.append({
                        "Employee ID": current_id,
                        "Name": record_name,
                        "Department": getattr(emp_match.department, 'name', 'N/A') if emp_match else "Unknown",
                        "Day": day_value,
                        "Time In": time_in if matched else None,
                        "Time Out": time_out if matched else None,
                        "Matched": matched
                    })

            db_employees = Employee.query.filter_by(active=True).all()
            excel_ids = {int(float(r["Employee ID"])) for r in records if r["Matched"]}

            for emp in db_employees:
                if emp.id not in excel_ids:
                    records.append({
                        "Employee ID": emp.id,
                        "Name": emp.get_full_name(),
                        "Department": getattr(emp.department, 'name', 'N/A') if emp.department else 'N/A',
                        "Day": attendance_date or datetime.now().date().isoformat(),
                        "Time In": None,
                        "Time Out": None,
                        "Matched": False
                    })

            if not records:
                flash("No valid attendance records found. Please check the Excel format.", "danger")
                return redirect(request.url)

            # Store in session
            session['import_attendance_preview'] = records
            preview_data = records
            flash("Preview loaded. Please confirm import.", "info")

        except Exception as e:
            flash(f"Error reading Excel file: {e}", "danger")
            return redirect(request.url)

    # Load preview from session if exists
    if 'import_attendance_preview' in session and not preview_data:
        preview_data = session['import_attendance_preview']

    return render_template('hr/admin/admin_import_attendance.html', preview=preview_data, employees=employees)


# ----------------- CONFIRM IMPORT -----------------
@hr_admin_bp.route('/add_attendance/confirm', methods=['POST'])
@login_required
def confirm_import_attendance():    
    import os
    records = session.get('import_attendance_preview', [])
    if not records:
        flash("No attendance records to import.", "danger")
        return redirect(url_for('hr_admin.add_attendance'))

    imported_count = 0

    for row in records:
        emp_id = row.get("Employee ID")
        try:
            emp_id_int = int(float(emp_id))
        except:
            continue

        emp = Employee.query.get(emp_id_int)
        if not emp:
            continue

        day = row.get("Day")
        if not day or "Tabling" in str(day):
            continue

        # Convert day string to date
        try:
            date_list = []
            if "~" in day:
                start_str, end_str = day.split("~")
                start_date = pd.to_datetime(start_str.strip(), errors='coerce').date()
                end_date = pd.to_datetime(end_str.strip(), errors='coerce').date()
                if start_date and end_date:
                    date_list = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]
            else:
                single_date = pd.to_datetime(day.strip(), errors='coerce').date()
                if single_date:
                    date_list = [single_date]
        except:
            continue

        time_in = row.get("Time In")
        time_out = row.get("Time Out")

        for att_date in date_list:
            # ✅ Skip if already exists
            existing = Attendance.query.filter_by(employee_id=emp.id, date=att_date).first()
            if existing:
                continue

            time_in_obj = pd.to_datetime(time_in, errors='coerce').time() if time_in else None
            time_out_obj = pd.to_datetime(time_out, errors='coerce').time() if time_out else None

            new_att = Attendance(
                employee_id=emp.id,
                date=att_date,
                time_in=time_in_obj,
                time_out=time_out_obj,
                status="Present" if time_in_obj else "Absent",
                remarks=""
            )
            db.session.add(new_att)
            imported_count += 1

    db.session.commit()
    session.pop('import_attendance_preview', None)

    # ✅ Cleanup uploaded files
    try:
        if os.path.exists(UPLOAD_FOLDER):
            for f in os.listdir(UPLOAD_FOLDER):
                os.remove(os.path.join(UPLOAD_FOLDER, f))
    except Exception as e:
        print(f"⚠️ Cleanup error: {e}")

    flash(f"✅ Successfully imported {imported_count} attendance record(s).", "success")
    return redirect(url_for('hr_admin.add_attendance'))

@hr_admin_bp.route('/add_manual_attendance', methods=['POST'])
@login_required
@admin_required  # Only admins can manually add attendance
def add_manual_attendance():
    employee_id = request.form.get('employee_id')
    date_str = request.form.get('date')
    time_in_str = request.form.get('time_in')
    time_out_str = request.form.get('time_out')
    status = request.form.get('status')
    remarks = request.form.get('remarks')

    # Validate employee
    emp = Employee.query.get(employee_id)
    if not emp:
        flash("Employee not found.", "danger")
        return redirect(url_for('hr_admin.add_attendance'))

    # Convert date and times
    try:
        attendance_date = datetime.strptime(date_str, '%Y-%m-%d').date()
    except:
        flash("Invalid date format.", "danger")
        return redirect(url_for('hr_admin.add_attendance'))

    time_in_obj = None
    time_out_obj = None
    try:
        if time_in_str:
            time_in_obj = datetime.strptime(time_in_str, '%H:%M').time()
        if time_out_str:
            time_out_obj = datetime.strptime(time_out_str, '%H:%M').time()
    except:
        flash("Invalid time format.", "danger")
        return redirect(url_for('hr_admin.add_attendance'))

    # Check if attendance already exists
    existing_att = Attendance.query.filter_by(employee_id=emp.id, date=attendance_date).first()
    if existing_att:
        flash(f"Attendance already exists for {emp.get_full_name()} on {attendance_date}.", "danger")
        return redirect(url_for('hr_admin.add_attendance'))

    # Create new attendance
    new_attendance = Attendance(
        employee_id=emp.id,
        date=attendance_date,
        time_in=time_in_obj,
        time_out=time_out_obj,
        status=status,
        remarks=remarks
    )

    db.session.add(new_attendance)
    db.session.commit()

    flash(f"Attendance for {emp.get_full_name()} on {attendance_date} added successfully.", "success")
    return redirect(url_for('hr_admin.add_attendance'))


@hr_admin_bp.route('/attendance/<int:attendance_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_attendance(attendance_id):
    attendance = Attendance.query.get_or_404(attendance_id)
    
    if request.method == 'POST':
        # Get form data
        time_in_str = request.form.get('time_in')  # e.g., "08:54"
        time_out_str = request.form.get('time_out')  # e.g., "17:00"
        status = request.form.get('modal_status')
        remarks = request.form.get('remarks', '')

        # Convert strings to datetime.time objects if not empty
        if time_in_str:
            h, m = map(int, time_in_str.split(":"))
            attendance.time_in = time(hour=h, minute=m)
        else:
            attendance.time_in = None

        if time_out_str:
            h, m = map(int, time_out_str.split(":"))
            attendance.time_out = time(hour=h, minute=m)
        else:
            attendance.time_out = None

        # Update other fields
        attendance.status = status
        attendance.remarks = remarks

        # Recalculate working hours
        attendance.calculate_working_hours()

        # Commit changes
        db.session.commit()
        flash('Attendance updated!', 'success')
        return redirect(url_for('hr_admin.attendance'))

    # GET request for modal JSON
    if request.headers.get('Accept') == 'application/json':
        return {
            "attendance_id": attendance.id,
            "date": attendance.date.strftime('%Y-%m-%d'),
            "time_in": attendance.time_in.strftime('%H:%M') if attendance.time_in else '',
            "time_out": attendance.time_out.strftime('%H:%M') if attendance.time_out else '',
            "modal_status": attendance.status,
            "remarks": attendance.remarks or ''
        }

    # fallback to page (optional)
    return render_template('hr/admin/edit_attendance_page.html', attendance=attendance)


    
@hr_admin_bp.route('/attendance/<int:attendance_id>/delete', methods=['POST'])
@login_required
@admin_required # Adjust as needed (e.g., @officer_required)
def delete_attendance(attendance_id):
    attendance_record = Attendance.query.get_or_404(attendance_id)
    try:
        db.session.delete(attendance_record)
        db.session.commit()
        flash('Attendance record deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting attendance record {attendance_id}: {e}")
        flash('Error deleting attendance record. Please try again.', 'error')

    return redirect(url_for('hr_admin.attendance'))


# ------------------------- Leaves -------------------------
@hr_admin_bp.route('/leaves')
@login_required
@admin_required
def view_leaves():
    page = request.args.get('page', 1, type=int)
    status_filter = request.args.get('status', '')

    query = Leave.query
    if status_filter:
        query = query.filter_by(status=status_filter)

    leaves = query.order_by(Leave.created_at.desc()).paginate(page=page, per_page=20, error_out=False)
    return render_template('hr/admin/admin_view_leaves.html', leaves=leaves, status_filter=status_filter)

"""
@hr_admin_bp.route('/review-leaves')
@login_required
@admin_required
def review_leaves():
    page = request.args.get('page', 1, type=int)
    status_filter = request.args.get('status', '')

    query = Leave.query

    # Apply filter if selected
    if status_filter:
        query = query.filter_by(status=status_filter)

    # ✅ Sort: Pending → Approved → Rejected → (None if any)
    query = query.order_by(
        db.case(
            (Leave.status == 'Pending', 0),
            (Leave.status == 'Approved', 1),
            (Leave.status == 'Rejected', 2),
            else_=3
        ),
        Leave.created_at.desc()
    )

    # Pagination
    leaves_paginated = query.paginate(page=page, per_page=20, error_out=False)
    
    return render_template(
        'hr/admin/review_leaves.html',
        leaves=leaves_paginated,
        status_filter=status_filter
    )
"""

"""
@hr_admin_bp.route('/leaves/<int:leave_id>/action', methods=['POST'])
@login_required
@admin_required
def leave_action(leave_id):
    leave = Leave.query.get_or_404(leave_id)
    action = request.form.get('action')  # 'Approved' or 'Rejected'
    comments = request.form.get('comments', '').strip()  # strip extra spaces

    leave.status = action
    leave.comments = comments if comments else None  # set None if empty
    leave.approved_by = current_user.id
    leave.approved_at = datetime.utcnow()

    try:
        db.session.commit()
        flash(f'Leave request {action.lower()} successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Leave action error: {e}")
        flash('Error updating leave request.', 'error')

    return redirect(url_for('hr_admin.view_leaves'))
"""


@hr_admin_bp.route('/departments')
@login_required
@admin_required
def view_departments():
    page = request.args.get('page', 1, type=int)
    search_query = request.args.get('search', '', type=str)

    query = Department.query

    if search_query:
        query = query.filter(Department.name.ilike(f'%{search_query}%'))

    departments = query.paginate(page=page, per_page=8)

    employee_counts = dict(
        db.session.query(
            Employee.department_id,
            func.count(Employee.id)
        ).group_by(Employee.department_id).all()
    )

    # If AJAX request, return only the cards HTML
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return render_template(
            'hr/admin/_department_cards.html', 
            departments=departments, 
            employee_counts=employee_counts
        )

    return render_template(
        'hr/admin/admin_view_departments.html',
        departments=departments,
        employee_counts=employee_counts
    )
@hr_admin_bp.route('/departments/<int:department_id>')
@login_required
@admin_required
def department_details(department_id):

    department = Department.query.get_or_404(department_id)

    # ✅ Force reload to avoid ORM stale relationship cache
    db.session.refresh(department)

    # ✅ Validate head belongs to department
    head = None
    if department.head_id:
        head = Employee.query.filter_by(
            id=department.head_id,
            department_id=department.id,
            status="Active"
        ).first()

    # ✅ Load department employees
    employees = Employee.query.filter(
        Employee.department_id == department_id,
        Employee.status == "Active"
    ).all()

    return render_template(
        'hr/admin/dept_details.html',
        department=department,
        employees=employees,
        head=head
    )

@hr_admin_bp.route('/departments/add', methods=['GET', 'POST'])
@login_required
@admin_required
def add_department():
    # Get employees who can be assigned as department head
    employees = Employee.query.join(User).filter(User.role.in_(['admin','officer','dept_head'])).all()

    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        head_id = request.form.get('head_id') or None

        department = Department(
            name=name,
            description=description,
            head_id=head_id
        )
        try:
            db.session.add(department)
            db.session.commit()

            # If a head is assigned, update their User role to 'dept_head'
            if head_id:
                employee = Employee.query.get(int(head_id))
                if employee and employee.user:
                    employee.user.role = 'dept_head'
                    db.session.commit()

            flash('Department added successfully!', 'success')
            return redirect(url_for('hr_admin.add_department'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error adding department: {str(e)}', 'error')

    return render_template('hr/admin/admin_add_dept.html', employees=employees)

@hr_admin_bp.route("/department/<int:department_id>/edit", methods=["POST"])
@login_required
@admin_required
def edit_department(department_id):

    department = Department.query.get_or_404(department_id)

    try:
        department.name = request.form.get("name") or department.name
        department.description = request.form.get("description") or department.description

        head_emp_id = request.form.get("dept_head")

        # ⭐ Get previous head
        previous_head = None
        if department.head_id:
            previous_head = Employee.query.get(department.head_id)

        # ⭐ If new head is selected
        if head_emp_id:

            new_head = Employee.query.get(int(head_emp_id))

            if not new_head:
                return jsonify(status="error", message="Employee not found.")

            if new_head.department_id != department.id:
                return jsonify(
                    status="error",
                    message="Employee must belong to this department."
                )

            # ✅ If head is changing → downgrade previous head
            if previous_head and previous_head.id != new_head.id:
                previous_head.role = "employee"

            # ✅ Assign new head
            new_head.role = "dept_head"
            department.head_id = new_head.id

        else:
            # ✅ Remove head if none selected
            if previous_head:
                previous_head.role = "employee"

            department.head_id = None

        db.session.commit()

        return jsonify(status="success", message="Department updated successfully!")

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(str(e))
        return jsonify(status="error", message="Error updating department.")

@hr_admin_bp.route('/hr/admin/positions')
@login_required
def view_positions():
    """Display paginated list of positions with employee count."""
    page = request.args.get('page', 1, type=int)
    per_page = 10

    positions = Position.query.order_by(Position.name.asc()).paginate(page=page, per_page=per_page)

    # Count employees in each position
    employee_counts = (
        db.session.query(Employee.position_id, db.func.count(Employee.id))
        .group_by(Employee.position_id)
        .all()
    )
    employee_counts = {pos_id: count for pos_id, count in employee_counts}

    return render_template(
        'hr/admin/admin_view_positions.html',
        positions=positions,
        employee_counts=employee_counts
    )


@hr_admin_bp.route("/hr/admin/add_position", methods=["GET", "POST"])
@login_required
@admin_required
def add_position():
    from main_app.models.hr_models import Position

    if request.method == "POST":
        name = request.form.get("name").strip()
        description = request.form.get("description").strip()

        if not name:
            flash("Position name is required.", "error")
            return redirect(url_for("hr_admin.add_position"))

        # Check for duplicate name
        existing_position = Position.query.filter_by(name=name).first()
        if existing_position:
            flash("A position with this name already exists.", "error")
            return redirect(url_for("hr_admin.add_position"))

        # Create and save new position
        new_position = Position(name=name, description=description)
        db.session.add(new_position)
        db.session.commit()

        flash(f"Position '{name}' added successfully!", "success")
        return redirect(url_for("hr_admin.view_positions"))  # You can adjust this target route

    return render_template("hr/admin/add_positions.html")

@hr_admin_bp.route("/position/<int:position_id>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def edit_position(position_id):
    position = Position.query.get_or_404(position_id)
    departments = Department.query.all()

    if request.method == "POST":
        try:
            position.name = request.form.get("name") or position.name
            position.description = request.form.get("description") or position.description
            dept_id = request.form.get("department_id")
            position.department_id = int(dept_id) if dept_id else position.department_id

            db.session.commit()

            # For AJAX request
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "success", "message": "Position updated successfully!"})

            flash("Position updated successfully!", "success")
            return redirect(url_for("hr_admin.view_positions"))

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating position {position_id}: {e}")

            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"status": "error", "message": "Error updating position. Please try again."})

            flash("Error updating position. Please try again.", "error")

    # Render form
    return render_template(
        "hr/admin/edit_position.html",
        position=position,
        departments=departments
    )

# ------------------------- Reports -------------------------
@hr_admin_bp.route('/reports')
@login_required
@admin_required
def reports():
    # Get filters from query parameters
    report_type = request.args.get('report_type', 'attendance')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    page = request.args.get('page', 1, type=int)

    # Convert dates to datetime objects if provided
    try:
        start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date() if start_date else None
        end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date() if end_date else None
    except ValueError:
        flash("Invalid date format", "error")
        start_date_obj = end_date_obj = None

    # Fetch data based on report type
    if report_type == 'attendance':
        query = Attendance.query
        if start_date_obj:
            query = query.filter(Attendance.date >= start_date_obj)
        if end_date_obj:
            query = query.filter(Attendance.date <= end_date_obj)
        data = query.order_by(Attendance.date.desc()).paginate(page=page, per_page=20)

        employees = Employee.query.filter_by(active=True).all()  # for filter dropdown

    elif report_type == 'leaves':
        query = Leave.query
        if start_date_obj:
            query = query.filter(Leave.start_date >= start_date_obj)
        if end_date_obj:
            query = query.filter(Leave.end_date <= end_date_obj)
        data = query.order_by(Leave.start_date.desc()).paginate(page=page, per_page=20)
        employees = Employee.query.filter_by(active=True).all()

    elif report_type == 'payroll':
        # Example payroll: just employees with salary (you can expand later)
        query = Employee.query.filter(Employee.salary != None)
        data = query.paginate(page=page, per_page=20)
        employees = None

    else:
        flash("Invalid report type", "error")
        return redirect(url_for('hr_admin.reports'))

    return render_template(
        'hr/admin/admin_reports.html',
        data=data,
        report_type=report_type,
        start_date=start_date or '',
        end_date=end_date or '',
        employees=employees if report_type in ['attendance', 'leaves'] else []
    )

@hr_admin_bp.route('/attendance-report', methods=['GET'])
@login_required
@admin_required
def attendance_report():
    # ------------------------------
    # Get filter params
    # ------------------------------
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    department_id = request.args.get('department_id')

    # Default to current month
    if not start_date:
        start_date = date.today().replace(day=1)
    else:
        start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
    
    if not end_date:
        end_date = date.today()
    else:
        end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

    # ------------------------------
    # Base employee query
    # ------------------------------
    employees_query = Employee.query.filter(Employee.status == "Active")
    if department_id:
        employees_query = employees_query.filter(Employee.department_id == department_id)
    employees = employees_query.all()
    total_employees = len(employees)

    # ------------------------------
    # Collect attendance data per employee
    # ------------------------------
    report_data = []
    total_hours_worked = 0
    total_present_days = 0

    for emp in employees:
        emp_attendances = Attendance.query.filter(
            Attendance.employee_id == emp.id,
            Attendance.date >= start_date,
            Attendance.date <= end_date
        ).all()

        days_present = sum(1 for a in emp_attendances if a.status in ["Present", "Late"])
        days_absent = sum(1 for a in emp_attendances if a.status == "Absent")
        late_count = sum(1 for a in emp_attendances if a.status == "Late")
        hours_worked = sum(a.working_hours for a in emp_attendances)

        total_hours_worked += hours_worked
        total_present_days += days_present

        report_data.append({
            "employee_name": emp.get_full_name(),
            "department_name": emp.department.name if emp.department else "",
            "days_present": days_present,
            "days_absent": days_absent,
            "late_count": late_count,
            "total_hours": round(hours_worked, 2)
        })

    # ------------------------------
    # Average attendance rate
    # ------------------------------
    avg_attendance_rate = (
        round((total_present_days / (total_employees * ((end_date - start_date).days + 1))) * 100, 2)
        if total_employees > 0 else 0
    )

    # ------------------------------
    # Department summary
    # ------------------------------
    department_summary = []
    departments = Department.query.all()
    for dept in departments:
        dept_emps = [e for e in employees if e.department_id == dept.id]
        if not dept_emps:
            continue

        dept_attendance_days = 0
        dept_total_hours = 0
        for emp in dept_emps:
            emp_att = [a for a in report_data if a["employee_name"] == emp.get_full_name()]
            if emp_att:
                dept_attendance_days += emp_att[0]["days_present"]
                dept_total_hours += emp_att[0]["total_hours"]
        
        num_days = ((end_date - start_date).days + 1) * len(dept_emps)
        avg_att = round((dept_attendance_days / num_days) * 100, 2) if num_days > 0 else 0
        avg_hours = round(dept_total_hours / len(dept_emps), 2) if dept_emps else 0

        department_summary.append({
            "name": dept.name,
            "avg_attendance": avg_att,
            "avg_hours": avg_hours
        })

    # ------------------------------
    # Render template
    # ------------------------------
    return render_template(
        "hr/admin/attendance_reports.html",
        report_data=report_data,
        department_summary=department_summary,
        total_employees=total_employees,
        total_hours_worked=round(total_hours_worked, 2),
        avg_attendance_rate=avg_attendance_rate,
        start_date=start_date,
        end_date=end_date,
        department_id=int(department_id) if department_id else "",
        departments=departments,
        current_date=date.today().strftime("%B %d, %Y"),
        current_user=current_user
    )



@hr_admin_bp.route('/attendance/reports/word')
@login_required
@admin_required
def attendance_report_word():
    # -----------------------------
    # Filters from GET
    # -----------------------------
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    department_id = request.args.get('department_id')

    # Default date range: last 30 days
    if start_date_str:
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
    else:
        start_date = date.today() - timedelta(days=30)

    if end_date_str:
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
    else:
        end_date = date.today()

    total_days = (end_date - start_date).days + 1

    # -----------------------------
    # Fetch employees (optional filter by department)
    # -----------------------------
    employees = Employee.query.filter(Employee.archived == False)
    if department_id:
        employees = employees.filter(Employee.department_id == department_id)
    employees = employees.all()

    # -----------------------------
    # Create Word Document
    # -----------------------------
    doc = Document()

    # Header
    header = doc.add_paragraph()
    header.alignment = 1  # center
    header.add_run("MUNICIPALITY OF NORZAGARAY\n").bold = True
    header.add_run("Attendance Report\n").bold = True
    header.add_run(f"From {start_date.strftime('%B %d, %Y')} to {end_date.strftime('%B %d, %Y')}\n").italic = True

    # Table header
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Employee Name"
    hdr_cells[1].text = "Department"
    hdr_cells[2].text = "Days Present"
    hdr_cells[3].text = "Days Absent"
    hdr_cells[4].text = "Total Hours Worked"

    # Attendance data
    for emp in employees:
        emp_att = Attendance.query.filter(
            Attendance.employee_id == emp.id,
            Attendance.date >= start_date,
            Attendance.date <= end_date
        ).all()

        days_present = sum(1 for a in emp_att if a.status in ["Present", "Late"])
        days_absent = sum(1 for a in emp_att if a.status == "Absent")
        total_hours = sum(a.working_hours for a in emp_att)

        row_cells = table.add_row().cells
        row_cells[0].text = emp.get_full_name()
        row_cells[1].text = emp.department.name if emp.department else "N/A"
        row_cells[2].text = str(days_present)
        row_cells[3].text = str(days_absent)
        row_cells[4].text = f"{total_hours:.2f}"

    # -----------------------------
    # Insights Section
    # -----------------------------
    doc.add_paragraph('\nOverall Insights', style='Heading 2')

    if employees:
        total_attendance_days = sum(
            sum(1 for a in Attendance.query.filter(
                Attendance.employee_id == emp.id,
                Attendance.date >= start_date,
                Attendance.date <= end_date
            ).all() if a.status in ["Present", "Late"]) for emp in employees
        )
        total_possible_days = total_days * len(employees)
        avg_attendance = round((total_attendance_days / total_possible_days) * 100, 2) if total_possible_days > 0 else 0

        total_working_hours = sum(
            sum(a.working_hours for a in Attendance.query.filter(
                Attendance.employee_id == emp.id,
                Attendance.date >= start_date,
                Attendance.date <= end_date
            ).all()) for emp in employees
        )
        avg_hours_per_employee = round(total_working_hours / len(employees), 2)

        doc.add_paragraph(f"Total Employees: {len(employees)}")
        doc.add_paragraph(f"Average Attendance: {avg_attendance}%")
        doc.add_paragraph(f"Average Hours Worked per Employee: {avg_hours_per_employee} hrs")

    # Department-wise insights
    doc.add_paragraph('\nDepartment-wise Insights', style='Heading 2')
    departments = Department.query.all()
    for dept in departments:
        dept_emps = [e for e in employees if e.department_id == dept.id]
        if not dept_emps:
            continue

        dept_attendance_days = 0
        dept_total_hours = 0

        for emp in dept_emps:
            emp_att = Attendance.query.filter(
                Attendance.employee_id == emp.id,
                Attendance.date >= start_date,
                Attendance.date <= end_date
            ).all()
            dept_attendance_days += sum(1 for a in emp_att if a.status in ["Present", "Late"])
            dept_total_hours += sum(a.working_hours for a in emp_att)

        num_days = total_days * len(dept_emps)
        avg_att = round((dept_attendance_days / num_days) * 100, 2) if num_days > 0 else 0
        avg_hours = round(dept_total_hours / len(dept_emps), 2) if dept_emps else 0

        doc.add_paragraph(f"{dept.name}: Avg Attendance: {avg_att}%, Avg Hours: {avg_hours}")

    # -----------------------------
    # Return as Word file
    # -----------------------------
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"Attendance_Report_{start_date}_{end_date}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@hr_admin_bp.route("/hr_admin/leave_report")
@admin_required
@login_required
def leave_report():
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    department_id = request.args.get('department_id')
    status_filter = request.args.get('status')

    start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date() if start_date_str else date.today() - timedelta(days=30)
    end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date() if end_date_str else date.today()

    employees = Employee.query.filter(Employee.archived == False)
    if department_id:
        employees = employees.filter(Employee.department_id == department_id)
    employees = employees.all()

    leave_data = []
    for emp in employees:
        emp_leaves = Leave.query.filter(
            Leave.employee_id == emp.id,
            Leave.start_date >= start_date,
            Leave.end_date <= end_date
        )
        if status_filter:
            emp_leaves = emp_leaves.filter(Leave.status == status_filter)
        leave_data.extend(emp_leaves.all())

    # Insights
    total_leaves = len(leave_data)
    avg_days_per_leave = round(sum(lv.days_requested for lv in leave_data)/total_leaves,2) if total_leaves else 0
    leave_types = [lv.leave_type.name for lv in leave_data if lv.leave_type]
    most_common_leave_type = Counter(leave_types).most_common(1)[0][0] if leave_types else "N/A"

    # Department-wise summary
    dept_summary = {}
    for dept in Department.query.all():
        dept_leaves = [lv for lv in leave_data if lv.employee.department_id == dept.id]
        if dept_leaves:
            dept_summary[dept.name] = {
                "total": len(dept_leaves),
                "avg_days": round(sum(lv.days_requested for lv in dept_leaves)/len(dept_leaves), 2)
            }

    return render_template(
        "hr/admin/leave_reports.html",
        leave_data=leave_data,
        start_date=start_date,
        end_date=end_date,
        departments=Department.query.all(),
        department_id=int(department_id) if department_id else None,
        total_leaves=total_leaves,
        avg_days_per_leave=avg_days_per_leave,
        most_common_leave_type=most_common_leave_type,
        dept_summary=dept_summary
    )

# ------------------------------
# Leave Report Word Export
# ------------------------------
@hr_admin_bp.route("/leave-report/word")
@login_required
@admin_required
def leave_report_word():
    # --- Filter params ---
    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")
    department_id = request.args.get("department_id", type=int)

    start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date() if start_date_str else None
    end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date() if end_date_str else None

    # --- Fetch leaves ---
    query = Leave.query.join(Employee).join(Department)
    if start_date:
        query = query.filter(Leave.start_date >= start_date)
    if end_date:
        query = query.filter(Leave.end_date <= end_date)
    if department_id:
        query = query.filter(Employee.department_id == department_id)

    all_leaves = query.order_by(Leave.start_date.asc()).all()

    # --- Create Word doc ---
    doc = Document()
    doc.add_heading("Municipality of Norzagaray", 0).alignment = 1
    doc.add_paragraph(f"Leave Report ({start_date_str or 'N/A'} - {end_date_str or 'N/A'})", style="Heading 1")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}\n")

    # --- Table ---
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    headers = ["Employee Name", "Department", "Leave Type", "Start Date", "End Date", "Days Requested", "Status"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for leave in all_leaves:
        row_cells = table.add_row().cells
        row_cells[0].text = leave.employee.get_full_name()
        row_cells[1].text = leave.employee.department.name if leave.employee.department else ""
        row_cells[2].text = leave.leave_type.name if leave.leave_type else ""
        row_cells[3].text = leave.start_date.strftime("%Y-%m-%d")
        row_cells[4].text = leave.end_date.strftime("%Y-%m-%d")
        row_cells[5].text = str(leave.days_requested)
        row_cells[6].text = leave.status

    # --- Insights ---
    doc.add_paragraph("\nInsights", style="Heading 2")
    total_leaves = len(all_leaves)
    avg_days = round(sum(lv.days_requested for lv in all_leaves)/total_leaves, 2) if total_leaves else 0
    doc.add_paragraph(f"Total leaves: {total_leaves}")
    doc.add_paragraph(f"Average leave days per record: {avg_days}")

    # --- Department-wise summary ---
    dept_summary = {}
    for dept in Department.query.all():
        dept_leaves = [lv for lv in all_leaves if lv.employee.department_id == dept.id]
        if dept_leaves:
            dept_summary[dept.name] = {
                "total": len(dept_leaves),
                "avg_days": round(sum(lv.days_requested for lv in dept_leaves) / len(dept_leaves), 2)
            }

    if dept_summary:
        doc.add_paragraph("\nDepartment-wise Summary", style="Heading 2")
        for dept, stats in dept_summary.items():
            doc.add_paragraph(f"{dept}: Total Leaves = {stats['total']}, Average Days = {stats['avg_days']}")

    # --- Send as file ---
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    filename = f"Leave_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@hr_admin_bp.route('/profile', methods=['GET'])
@login_required
@admin_required
def profile():
    user = current_user
    employee = user.employee_profile


    age = None
    working_duration = None
    if employee:
        if employee.date_of_birth:
            today = date.today()
            age = today.year - employee.date_of_birth.year - ((today.month, today.day) < (employee.date_of_birth.month, employee.date_of_birth.day))
        working_duration = employee.get_working_duration()


    return render_template(
    "hr/admin/profile.html",
    user=user,
    employee=employee,
    age=age,
    working_duration=working_duration
    )


@hr_admin_bp.route('/profile/edit', methods=['POST'])
@login_required
@admin_required
def edit_profile():
    user = current_user
    employee = user.employee_profile


    data = request.get_json()
    current_password = data.get('current_password')
    new_email = data.get('email')
    new_password = data.get('new_password')
    confirm_password = data.get('confirm_password')


    # Verify current password
    if current_password != user.password:
        return jsonify({'status': 'error', 'message': 'Current password is incorrect.'}), 400


    # Update email
    if new_email and new_email != user.email:
        existing_user = User.query.filter_by(email=new_email).first()
        if existing_user:
            return jsonify({'status': 'error', 'message': 'Email already in use.'}), 400
        user.email = new_email
        if employee:
            employee.email = new_email


    # Update password
    if new_password:
        if new_password != confirm_password:
            return jsonify({'status': 'error', 'message': 'Passwords do not match.'}), 400
        user.password = new_password # plain text (for now)


    db.session.commit()
    return jsonify({'status': 'success', 'message': 'Profile updated successfully.'})